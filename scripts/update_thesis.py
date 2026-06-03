import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'C:\Users\musha\Desktop\QuPath\thesis_final.docx')

def set_para_text(para, text):
    """Replace all runs in a paragraph with a single run containing text."""
    p_elem = para._p
    # Remove all run elements
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    # Add new run
    run = para.add_run(text)
    return run

def insert_paragraph_after(ref_para, text):
    """Insert a new paragraph after ref_para with same paragraph properties."""
    new_p = OxmlElement('w:p')
    # Copy paragraph properties from ref_para if any
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_pPr = copy.deepcopy(pPr)
        new_p.append(new_pPr)
    # Add run with text
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    # Insert after ref_para
    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())

# ============================================================
# UPDATE 1: Abstract para 13
# ============================================================
para13 = doc.paragraphs[13]
set_para_text(para13,
    "A digital pathology workflow implemented in QuPath enabled objective, reproducible quantification of "
    "tumour cell density and lymphocyte density within manually annotated tumour regions of interest across "
    "141 cases with complete slide-to-clinical-record linkage. Quantitative metrics demonstrated wide "
    "biological variability across the cohort (lymphocyte density: median 101.9 cells/mm2, range "
    "0-3478.8 cells/mm2; tumour cell density: median 1112.0 cells/mm2, range 4.0-3792.9 cells/mm2) "
    "and excellent intra-tumoral reproducibility (inter-ROI Pearson r = 0.882 for lymphocyte density and "
    "r = 0.907 for tumour cell density; both p < 0.001). However, in this cohort, digitally quantified "
    "lymphocyte density and tumour cell density did not demonstrate statistically significant associations "
    "with overall survival, STAS status, nodal involvement, or pathologic stage, likely reflecting "
    "methodological limitations of the threshold-based cell classification approach and the moderate "
    "statistical power of the current series. These findings provide a validated analytical framework and "
    "highlight the methodological refinements -- including deep learning-based cell classification and "
    "spatial immune analysis -- needed to fully harness the prognostic potential of quantitative digital "
    "pathology in surgically resected lung adenocarcinoma."
)
print("Updated para 13 (Abstract)")

# ============================================================
# UPDATE 2: Figure 4.8 caption (para 242) - remove pending language
# ============================================================
para242 = doc.paragraphs[242]
set_para_text(para242,
    "Figure 4.8 presents Kaplan-Meier overall survival curves for each of the three primary quantitative "
    "metrics, dichotomised at the cohort median. The curves stratify survival by high versus low lymphocyte "
    "density (median split 109.2 cells/mm2; log-rank p = 0.855), high versus low tumour cell density "
    "(median split 1122.4 cells/mm2; log-rank p = 0.530), and high versus low tumour-to-lymphocyte ratio "
    "(median split 6.2; log-rank p = 0.596). The curves for all three metrics are closely overlapping, "
    "consistent with the absence of statistically significant survival differences between high and low "
    "subgroups for any of the quantitative digital pathology variables."
)
print("Updated para 242 (Figure 4.8 caption)")

# ============================================================
# UPDATE 3: Insert two new paragraphs after para 248
# ============================================================
para248 = doc.paragraphs[248]

null_assoc_text = (
    "The systematic analysis of associations between digitally quantified metrics and conventional "
    "clinicopathologic features -- including STAS, nodal status, pathologic stage, lymphovascular "
    "invasion, and visceral pleural invasion -- revealed no statistically significant relationships for "
    "any of the three primary quantitative variables (all Mann-Whitney U p > 0.30; Section 4.6). This "
    "absence of association with conventional adverse histopathologic features warrants careful "
    "interpretation. The most probable explanation lies in the methodological limitations of the "
    "threshold-based cell classification strategy: the subtractive definition of tumour cells as all "
    "detected nuclei not meeting morphological lymphocyte criteria is susceptible to misclassification of "
    "stromal fibroblasts, macrophages, and plasma cells as neoplastic cells, particularly in cases with "
    "dense desmoplastic stroma or prominent inflammatory infiltration. Such misclassification would "
    "introduce random noise into the density estimates, attenuating any underlying biological signal and "
    "biasing findings toward the null. A complementary biological interpretation is that, in lung "
    "adenocarcinoma, global density of immune infiltration and neoplastic cells within representative "
    "tumour ROIs may not be strongly mechanistically coupled to the histopathologic features of local "
    "invasion and nodal metastasis, which are governed more by epithelial-to-mesenchymal transition "
    "mechanisms than by immune surveillance dynamics."
)

null_survival_text = (
    "Kaplan-Meier survival analysis with median dichotomisation of each metric demonstrated no significant "
    "difference in overall survival between high and low subgroups for lymphocyte density (log-rank "
    "chi2 = 0.033, p = 0.855), tumour cell density (log-rank chi2 = 0.394, p = 0.530), or "
    "the tumour-to-lymphocyte ratio (log-rank chi2 = 0.280, p = 0.596). Continuous-variable Cox "
    "proportional hazards regression likewise yielded no significant prognostic associations "
    "(lymphocyte density: HR = 1.03, 95% CI 0.91-1.16, p = 0.661; tumour cell density: HR = 1.04, "
    "95% CI 0.85-1.27, p = 0.713; tumour:lymphocyte ratio: HR = 0.96, 95% CI 0.78-1.16, "
    "p = 0.651). The multivariable extended model incorporating digital pathology variables alongside "
    "conventional prognostic factors produced a marginally lower C-index than the base clinical model "
    "alone (0.631 versus 0.640), confirming that the digital metrics provided no incremental prognostic "
    "discrimination in this cohort. These results contrast with studies demonstrating prognostic "
    "significance of computationally quantified immune infiltration in non-small cell lung cancer "
    "(Brambilla et al., 2016; Corredor et al., 2019); however, those studies employed immunohistochemical "
    "CD8+ T-cell quantification or deep learning-based cell classification rather than threshold-based "
    "H&E classification -- methodological differences that may substantially account for the "
    "discrepancy in findings. With 63 events across 141 analysed cases, the present study had sufficient "
    "power to detect hazard ratios of approximately 1.5 or greater but was underpowered to detect "
    "more modest effects that could nonetheless be clinically meaningful."
)

new_para_assoc = insert_paragraph_after(para248, null_assoc_text)
print("Inserted null association paragraph after para 248")

new_para_surv = insert_paragraph_after(new_para_assoc, null_survival_text)
print("Inserted null survival paragraph after new assoc paragraph")

# ============================================================
# UPDATE 4: Limitations paragraph - remove "pending" last sentence
# ============================================================
lim_para = None
for p in doc.paragraphs:
    if p.text.startswith("Several limitations of the present analytical approach"):
        lim_para = p
        break

if lim_para:
    old_text = lim_para.text
    pending_start = "Fourth, the slide-to-case data integration is currently pending"
    if pending_start in old_text:
        cut_idx = old_text.rfind(pending_start)
        new_text = old_text[:cut_idx].rstrip()
        new_text = (new_text +
            " Fourth, with 63 events across 141 analysed cases, statistical power was sufficient to "
            "detect hazard ratios of approximately 1.5 or greater but was limited for detecting more "
            "modest effect sizes; the null survival findings should therefore be interpreted as "
            "inconclusive rather than as definitive evidence against the prognostic utility of "
            "quantitative digital morphology in this disease context."
        )
        set_para_text(lim_para, new_text)
        print("Updated limitations paragraph (removed pending)")
    else:
        print("WARNING: pending text not found in limitations paragraph")
else:
    print("WARNING: Limitations paragraph not found")

# ============================================================
# UPDATE 5: 4.11 Conclusion - "pending integration" paragraph
# ============================================================
integ_para = None
for p in doc.paragraphs:
    if p.text.startswith("The integration of these quantitative metrics with the clinical outcome data"):
        integ_para = p
        break

if integ_para:
    set_para_text(integ_para,
        "The integration of these quantitative metrics with clinical outcome data -- achieved through "
        "systematic slide-to-case linkage for 141 of 173 cases, Kaplan-Meier survival analysis, and "
        "multivariable Cox proportional hazards regression -- revealed that, in this cohort, "
        "lymphocyte density, tumour cell density, and the tumour-to-lymphocyte ratio did not demonstrate "
        "statistically significant associations with overall survival or with established adverse "
        "clinicopathologic features. While these null findings do not support the prognostic stratification "
        "hypothesis in the current implementation, they provide an important empirical baseline and "
        "highlight the methodological refinements -- principally deep learning-based cell "
        "classification and spatial immune analysis -- required to fully characterise the prognostic "
        "potential of computationally derived morphological metrics in surgically resected lung "
        "adenocarcinoma. The validated QuPath workflow, the annotated case cohort, and the analytical "
        "framework established in this chapter provide the foundation for such prospective refinement."
    )
    print("Updated 4.11 conclusion paragraph")
else:
    print("WARNING: 4.11 pending paragraph not found")

# ============================================================
# UPDATE 6: 5.5 Clinical Implications, first paragraph
# ============================================================
clin_imp_para = None
for p in doc.paragraphs:
    if p.text.startswith("The clinical implications of this work are contingent"):
        clin_imp_para = p
        break

if clin_imp_para:
    set_para_text(clin_imp_para,
        "The clinical implications of this thesis must be interpreted in light of the actual outcome of "
        "the survival and association analyses presented in Chapter 4. Systematic analysis of 141 cases "
        "with linked clinical and digital pathology data revealed that computationally quantified "
        "lymphocyte density, tumour cell density, and the tumour-to-lymphocyte ratio did not demonstrate "
        "statistically significant associations with overall survival, STAS status, nodal involvement, "
        "pathologic stage, lymphovascular invasion, or visceral pleural invasion. These null findings "
        "carry important scientific meaning. They indicate that, in the present threshold-based "
        "implementation, the quantitative metrics do not generate sufficient biological signal to "
        "stratify patients by clinicopathologic or prognostic subgroup in this cohort. They do not, "
        "however, negate the established prognostic framework from Chapter 2 -- STAS, VPI, and nodal "
        "burden remain the evidence-based determinants of risk in surgically resected lung "
        "adenocarcinoma and should continue to anchor prognostic reporting in routine practice."
    )
    print("Updated 5.5 Clinical Implications para 1")
else:
    print("WARNING: 5.5 para 1 not found")

# ============================================================
# UPDATE 7: 5.5 second paragraph
# ============================================================
stas_corr_para = None
for p in doc.paragraphs:
    if p.text.startswith("If tumour cell density correlates with pathological stage"):
        stas_corr_para = p
        break

if stas_corr_para:
    set_para_text(stas_corr_para,
        "The null prognostic findings do not preclude the future utility of digital pathology "
        "approaches in this clinical context. Several methodological factors likely contributed to the "
        "absence of prognostic signal. The subtractive definition of tumour cells -- classifying all "
        "non-lymphocyte-like nuclei as neoplastic -- is inherently imprecise and probably introduced "
        "misclassification of stromal fibroblasts, macrophages, and endothelial cells that attenuated "
        "any underlying biological association. The restriction to two four-millimetre-square ROIs per "
        "case, while validated by high inter-ROI correlations, does not capture spatial heterogeneity "
        "at the tumour invasion front or within tertiary lymphoid structures -- compartments shown in "
        "recent studies to carry greater prognostic information than the intratumoral average. "
        "Additionally, with 63 events across 141 cases and a follow-up window bounded by the "
        "2015-2019 surgical period, the study had limited power to detect modest effect sizes. The "
        "workflow validated here provides the methodological infrastructure for future implementations "
        "using deep learning-based cell classification trained on expert-annotated cell-level ground "
        "truth, which would substantially reduce misclassification error and potentially unmask "
        "prognostic associations that the present threshold-based approach was unable to detect."
    )
    print("Updated 5.5 Clinical Implications para 2")
else:
    print("WARNING: 5.5 para 2 not found")

# ============================================================
# UPDATE 8: 5.5 third paragraph
# ============================================================
health_sys_para = None
for p in doc.paragraphs:
    if p.text.startswith("From a health system perspective, the most compelling feature"):
        health_sys_para = p
        break

if health_sys_para:
    set_para_text(health_sys_para,
        "From a health system perspective, the validated digital pathology workflow described in this "
        "thesis -- despite yielding null prognostic findings in the current application -- "
        "provides a reproducible, freely available analytical infrastructure applicable to future "
        "validation cohorts without additional laboratory expenditure. The quantitative metrics are "
        "derived from existing H&E slides using open-source software, requiring only whole slide "
        "scanning infrastructure that many academic pathology departments have already acquired. "
        "Should future implementations using deep learning-based cell classification, spatial immune "
        "analysis, or larger multi-centre cohorts demonstrate prognostic associations, the analytical "
        "framework described in this thesis would provide the methodological template for clinical "
        "translation. The present null results should be communicated as methodologically informative "
        "negative findings: they demonstrate the feasibility and reproducibility of the workflow at "
        "scale, establish a quantitative baseline for the cohort, and precisely identify the "
        "methodological bottlenecks -- cell classification precision and spatial resolution -- "
        "that must be addressed before digital morphological metrics can contribute to clinical "
        "prognostication in surgically resected lung adenocarcinoma."
    )
    print("Updated 5.5 Clinical Implications para 3")
else:
    print("WARNING: 5.5 para 3 not found")

# ============================================================
# UPDATE 9: 5.6 Limitations - remove "currently pending" sentence
# ============================================================
spat_lim_para = None
for p in doc.paragraphs:
    if p.text.startswith("The restriction of the digital analysis to two regions of interest per case"):
        spat_lim_para = p
        break

if spat_lim_para:
    old_text = spat_lim_para.text
    pending_phrase = "The case-level linkage between slide identifiers and clinical records required careful validation and is currently pending completion"
    if pending_phrase in old_text:
        cut_idx = old_text.rfind(pending_phrase)
        new_text = old_text[:cut_idx].rstrip().rstrip(';')
        new_text = (new_text +
            ". Case-level linkage between slide identifiers and clinical records was achieved for "
            "141 of 173 cases (81.5%), with 32 cases failing to map due to incomplete or ambiguous "
            "slide identifier data. This partial linkage rate introduces a small possibility of "
            "selection bias in the analytical cohort, as unmapped cases may differ systematically "
            "from mapped cases in terms of clinicopathologic characteristics."
        )
        set_para_text(spat_lim_para, new_text)
        print("Updated 5.6 Limitations para (removed pending)")
    else:
        print("WARNING: pending phrase not found in 5.6 limitations para")
        print("Last 400 chars:", old_text[-400:])
else:
    print("WARNING: 5.6 limitations para not found")

# ============================================================
# UPDATE 10: Chapter 6 last conclusion paragraph
# ============================================================
prog_fw_para = None
for p in doc.paragraphs:
    if ("progressive conceptual framework" in p.text and
        ("currently in progress" in p.text or "will determine whether" in p.text)):
        prog_fw_para = p
        break

if not prog_fw_para:
    # Try broader search
    for p in doc.paragraphs:
        if p.text.startswith("The progressive conceptual framework of this thesis"):
            prog_fw_para = p
            break

if prog_fw_para:
    set_para_text(prog_fw_para,
        "The progressive conceptual framework of this thesis -- from established clinicopathologic "
        "factors, through under-recognised histologic findings, to quantitative digital morphology -- "
        "illustrates the untapped potential of routine haematoxylin and eosin histology as a source of "
        "prognostically relevant information. In the present implementation, digitally quantified "
        "morphological metrics did not demonstrate independent prognostic value beyond conventional "
        "clinicopathologic factors in the study cohort, a finding that underscores the importance of "
        "methodological refinement -- particularly the transition from threshold-based to deep "
        "learning-based cell classification -- before clinical translation can be pursued. The "
        "validated QuPath workflow, the annotated whole slide image dataset, and the analytical "
        "infrastructure established in this thesis provide the methodological foundation for prospective "
        "multi-centre studies with the statistical power to definitively address the prognostic utility "
        "of quantitative digital morphology. As pathology advances toward a more quantitative and "
        "data-driven discipline, the systematic extraction of cellular and spatial metrics from "
        "digitised slides using open-source computational tools represents an accessible, reproducible "
        "step toward improved tumour characterisation and risk stratification in surgically resected "
        "lung adenocarcinoma."
    )
    print("Updated Chapter 6 last conclusion paragraph")
else:
    print("WARNING: Chapter 6 last conclusion paragraph not found")

# Save
doc.save(r'C:\Users\musha\Desktop\QuPath\thesis_final.docx')
print("\nAll updates complete. Saved thesis_final.docx")
