#!/usr/bin/env python3
"""
Update thesis_draft.docx with actual statistical results.
"""

import json
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import re

# ─── Load all results ─────────────────────────────────────────────────────────
with open('results_cache.json') as f:
    R = json.load(f)

with open('km_results_correct.json') as f:
    KM = json.load(f)

merged = pd.read_csv('analysis_results.csv')
complete = merged.dropna(subset=['lymp_mean','tumor_mean','lymp_ROI1','lymp_ROI2',
                                   'tumor_ROI1','tumor_ROI2']).copy()

# ─── Helper ───────────────────────────────────────────────────────────────────
def fp(p):
    if p < 0.001: return "p<0.001"
    return f"p={p:.3f}"

def sig_phrase(p):
    if p < 0.05:   return "statistically significant"
    if p < 0.10:   return "a trend toward significance"
    return "no statistically significant difference"

def hr_str(hr, lo, hi, p):
    return f"HR = {hr:.2f} (95% CI {lo:.2f}–{hi:.2f}), {fp(p)}"

def iqr_str(med, q1, q3):
    return f"{med:.1f} ({q1:.1f}–{q3:.1f})"

# ─── Build replacement text for each section ──────────────────────────────────

# ── 4.4 Quantitative Variables ────────────────────────────────────────────────
d_lymp  = R['desc']['lymp_mean']
d_tumor = R['desc']['tumor_mean']
d_ratio = R['desc']['ratio_mean']
n       = R['n_complete']

sec44_paras = [
    (
        f"The primary quantitative output of the digital pathology workflow described in "
        f"Section 4.3 consisted of three inter-related variables quantified per case: "
        f"per-case mean lymphocyte density, per-case mean tumour cell density, and the "
        f"tumour-to-lymphocyte ratio. These were derived from two rectangular ROIs "
        f"(each 4 mm²) annotated per slide, as described above. "
        f"The analytical cohort comprised {n} cases with complete density data."
    ),
    (
        f"Lymphocyte density across the {n} analysed cases demonstrated a highly right-skewed "
        f"distribution, with a mean of {d_lymp['mean']:.1f} cells/mm² (SD "
        f"{d_lymp['sd']:.1f} cells/mm²), a median of {d_lymp['median']:.1f} "
        f"cells/mm² (IQR {d_lymp['q1']:.1f}–{d_lymp['q3']:.1f} cells/mm²), "
        f"and a range of {d_lymp['min']:.1f}–{d_lymp['max']:.1f} cells/mm². "
        f"This distributional pattern is consistent with the well-documented heterogeneity of "
        f"tumour-infiltrating lymphocyte density in lung adenocarcinoma."
    ),
    (
        f"Tumour cell density was more symmetrically distributed, with a mean of "
        f"{d_tumor['mean']:.1f} cells/mm² (SD {d_tumor['sd']:.1f} cells/mm²), "
        f"a median of {d_tumor['median']:.1f} cells/mm² (IQR {d_tumor['q1']:.1f}–"
        f"{d_tumor['q3']:.1f} cells/mm²), and a range of {d_tumor['min']:.1f}–"
        f"{d_tumor['max']:.1f} cells/mm²."
    ),
    (
        f"The tumour-to-lymphocyte ratio (tumour cell density divided by lymphocyte density) "
        f"was calculable in {d_ratio['n']} cases and also exhibited right skew, with a mean "
        f"of {d_ratio['mean']:.1f} (SD {d_ratio['sd']:.1f}), a median of "
        f"{d_ratio['median']:.1f} (IQR {d_ratio['q1']:.1f}–{d_ratio['q3']:.1f}), "
        f"and a range of {d_ratio['min']:.1f}–{d_ratio['max']:.1f}. "
        f"Given the skewed distributions of all three variables, non-parametric methods "
        f"were used for all group comparisons, and natural logarithm transformation was applied "
        f"prior to Cox regression analyses."
    ),
]

# ── 4.5 Intra-tumoral Variability ─────────────────────────────────────────────
iv_lymp  = R['ivar']['lymp']
iv_tumor = R['ivar']['tumor']

sec45_paras = [
    (
        f"Intra-tumoral heterogeneity — the spatial variability of cellular composition "
        f"within a single neoplasm — is a fundamental biological property of solid tumours "
        f"that carries direct implications for any sampling-based quantification strategy. "
        f"The two-ROI per case design employed in this study was evaluated for its ability to "
        f"provide reproducible, representative measurements of both lymphocyte and tumour cell "
        f"density across the tumour parenchyma."
    ),
    (
        f"Across the full analytical cohort of {n} cases, {n * 2} ROIs were analysed in total "
        f"(two per case, each 4 mm²). Lymphocyte density exhibited a highly right-skewed "
        f"distribution, as detailed in Section 4.4. The coefficient of variation (CV) across "
        f"all individual ROIs was {iv_lymp['cv']:.1f}%, reflecting substantial between-case "
        f"heterogeneity in the immune microenvironment."
    ),
    (
        f"Tumour cell density was more uniformly distributed across the {n * 2} ROIs, with a CV "
        f"of {iv_tumor['cv']:.1f}% (mean {d_tumor['mean']:.1f} cells/mm², "
        f"SD {d_tumor['sd']:.1f} cells/mm²)."
    ),
    (
        f"The degree of intra-tumoral variability was assessed by computing the Pearson "
        f"correlation coefficient between ROI 1 and ROI 2 measurements across all "
        f"{n} cases, as well as the coefficient of variation (CV) across the two ROIs per case. "
        f"For lymphocyte density, the inter-ROI Pearson correlation was "
        f"r = {iv_lymp['r']:.3f} ({fp(iv_lymp['p'])}), indicating a strong linear "
        f"relationship between measurements obtained from the two sampling regions within each tumour."
    ),
    (
        f"For tumour cell density, the inter-ROI Pearson correlation was "
        f"r = {iv_tumor['r']:.3f} ({fp(iv_tumor['p'])}), indicating excellent "
        f"reproducibility of tumour cellularity measurements between the two sampling regions."
    ),
    (
        f"The strong between-ROI correlations observed for both lymphocyte density "
        f"(r = {iv_lymp['r']:.3f}) and tumour cell density (r = {iv_tumor['r']:.3f}) "
        f"provide important methodological validation for the two-ROI sampling strategy employed "
        f"in this study, and support the use of per-case mean values as reliable summary measures "
        f"for subsequent statistical analyses."
    ),
    (
        f"In summary, the intra-tumoral variability analysis demonstrates that the two-ROI per "
        f"case sampling strategy captures biologically meaningful and reproducible information "
        f"about both the immune and neoplastic cellular compartments of lung adenocarcinoma, "
        f"supporting the use of mean values derived from two ROIs as reliable summary measures "
        f"for each case."
    ),
]

# ── 4.6 Association with Clinicopathologic Features ───────────────────────────
# Retrieve all association results
assoc = R['assoc']
comps = R['comparisons']

def assoc_sentence(label, col, name0, name1, metric_key, metric_label, unit='cells/mm²'):
    key = f"{label}_{metric_key}"
    d   = assoc[key]
    med0 = d['group0_median']; q10 = d['group0_q1']; q30 = d['group0_q3']
    med1 = d['group1_median']; q11 = d['group1_q1']; q31 = d['group1_q3']
    n0 = d['group0_n']; n1 = d['group1_n']
    U = d['U']; p = d['p']
    p_str = fp(p)
    direction = ''
    if p < 0.05:
        if med0 > med1:
            direction = f'{name0} cases had significantly higher {metric_label} than {name1} cases'
        else:
            direction = f'{name1} cases had significantly higher {metric_label} than {name0} cases'
        return (f"{direction} (median {iqr_str(med0, q10, q30)} vs "
                f"{iqr_str(med1, q11, q31)} {unit}; U = {U:.0f}, {p_str}).")
    elif p < 0.10:
        return (f"There was a trend toward higher {metric_label} in {name0 if med0>med1 else name1} "
                f"cases compared with {name1 if med0>med1 else name0} cases "
                f"(median {iqr_str(med0, q10, q30)} vs {iqr_str(med1, q11, q31)} {unit}; "
                f"U = {U:.0f}, {p_str}), although this did not reach statistical significance.")
    else:
        return (f"No statistically significant difference in {metric_label} was observed between "
                f"{name0} and {name1} cases (median {iqr_str(med0, q10, q30)} vs "
                f"{iqr_str(med1, q11, q31)} {unit}; U = {U:.0f}, {p_str}).")

sec46_paras = [
    (
        f"The biological interpretation of lymphocyte density and tumour cell density as "
        f"prognostically relevant metrics depends, in part, on whether these quantitative "
        f"variables exhibit systematic associations with established clinicopathological "
        f"determinants of tumour biology and prognosis. Associations between the three "
        f"primary quantitative variables — per-case mean lymphocyte density, per-case "
        f"mean tumour cell density, and the tumour-to-lymphocyte ratio — and five "
        f"clinicopathological features were evaluated using the Mann-Whitney U test, "
        f"with results presented as median (IQR)."
    ),
    # STAS
    (
        f"Spread through air spaces (STAS) was present in 36 of {n} cases (25.5%). "
        + assoc_sentence('STAS','STAS','STAS-negative','STAS-positive','lymp_mean','lymphocyte density')
        + " " + assoc_sentence('STAS','STAS','STAS-negative','STAS-positive','tumor_mean','tumour cell density')
        + " " + assoc_sentence('STAS','STAS','STAS-negative','STAS-positive','ratio_mean','tumour-to-lymphocyte ratio',unit='')
    ),
    # pN
    (
        f"Nodal metastasis (pN1 or pN2) was present in 35 of {n} cases (24.8%). "
        + assoc_sentence('pN','pN','N0','N+','lymp_mean','lymphocyte density')
        + " " + assoc_sentence('pN','pN','N0','N+','tumor_mean','tumour cell density')
        + " " + assoc_sentence('pN','pN','N0','N+','ratio_mean','tumour-to-lymphocyte ratio',unit='')
    ),
    # Stage
    (
        f"Stage grouping (Stage I: groups 1–4 corresponding to IA1, IA2, IA3, IB; "
        f"Stage II/III: groups 5–8) showed no significant association with any of "
        f"the three quantitative metrics. "
        + assoc_sentence('Stage','Stage_I_vs_higher','Stage I','Stage II/III','lymp_mean','lymphocyte density')
        + " " + assoc_sentence('Stage','Stage_I_vs_higher','Stage I','Stage II/III','tumor_mean','tumour cell density')
        + " " + assoc_sentence('Stage','Stage_I_vs_higher','Stage I','Stage II/III','ratio_mean','tumour-to-lymphocyte ratio',unit='')
    ),
    # LVI
    (
        f"Lymphovascular invasion (LVI) was present in 43 of {n} cases (30.5%). "
        + assoc_sentence('LVI','Lymphovascular_Invasion','LVI-negative','LVI-positive','lymp_mean','lymphocyte density')
        + " " + assoc_sentence('LVI','Lymphovascular_Invasion','LVI-negative','LVI-positive','tumor_mean','tumour cell density')
        + " " + assoc_sentence('LVI','Lymphovascular_Invasion','LVI-negative','LVI-positive','ratio_mean','tumour-to-lymphocyte ratio',unit='')
    ),
    # VPI
    (
        f"Visceral pleural invasion (VPI) was present in 81 of {n} cases (57.4%). "
        + assoc_sentence('VPI','Visceral_Pleural_Invasion','VPI-negative','VPI-positive','lymp_mean','lymphocyte density')
        + " " + assoc_sentence('VPI','Visceral_Pleural_Invasion','VPI-negative','VPI-positive','tumor_mean','tumour cell density')
        + " " + assoc_sentence('VPI','Visceral_Pleural_Invasion','VPI-negative','VPI-positive','ratio_mean','tumour-to-lymphocyte ratio',unit='')
    ),
    (
        f"In summary, none of the three quantitative digital pathology variables — "
        f"lymphocyte density, tumour cell density, or their ratio — demonstrated "
        f"statistically significant associations with any of the five clinicopathological "
        f"features examined (STAS, nodal status, tumour stage, lymphovascular invasion, or "
        f"visceral pleural invasion) at the p<0.05 threshold. This absence of significant "
        f"associations is discussed in the context of the study’s statistical power and "
        f"the biological heterogeneity of the cohort in Section 4.10."
    ),
]

# ── 4.7 Survival Analysis ─────────────────────────────────────────────────────
n_surv = R['n_surv']
n_ev   = R['n_events']

km_lymp  = KM['lymp_mean']
km_tumor = KM['tumor_mean']
km_ratio = KM['ratio_mean']
cox      = R['cox_univar']

def km_sentence(d, metric_label, higher_is_better=True):
    thr = d['threshold']
    med_low  = f"{d['med_low']:.1f} months"  if d['med_low']  else "not reached"
    med_high = f"{d['med_high']:.1f} months" if d['med_high'] else "not reached"
    p_str = fp(d['logrank_p'])
    os5_low  = d['os5_low']  * 100
    os5_high = d['os5_high'] * 100
    sig = ''
    if d['logrank_p'] < 0.05:
        sig = (f"The difference in survival between groups was statistically significant "
               f"(log-rank {p_str}).")
    elif d['logrank_p'] < 0.10:
        sig = f"The difference showed a trend toward significance (log-rank {p_str})."
    else:
        sig = f"The difference in survival between groups was not statistically significant (log-rank {p_str})."
    return (
        f"For {metric_label}, cases were dichotomised at the median value of {thr:.1f} into low "
        f"(n = {d['n_low']}, events = {d['events_low']}) and high "
        f"(n = {d['n_high']}, events = {d['events_high']}) subgroups. "
        f"The five-year overall survival estimate was {os5_low:.1f}% in the low group and "
        f"{os5_high:.1f}% in the high group; median overall survival was {med_low} and "
        f"{med_high}, respectively. {sig}"
    )

sec47_paras = [
    (
        f"The prognostic relevance of digitally quantified lymphocyte density, tumour cell "
        f"density, and the tumour-to-lymphocyte ratio with respect to overall survival was "
        f"investigated in {n_surv} cases with complete density and survival data, "
        f"comprising {n_ev} death events and {n_surv - n_ev} censored observations. "
        f"Median follow-up in the analytical cohort was 71.5 months (IQR "
        f"30.8–88.0 months; range 0–114 months)."
    ),
    km_sentence(km_lymp, 'lymphocyte density'),
    km_sentence(km_tumor, 'tumour cell density', higher_is_better=False),
    km_sentence(km_ratio, 'the tumour-to-lymphocyte ratio', higher_is_better=False),
    (
        f"Univariate Cox proportional-hazards regression was performed for each quantitative "
        f"variable entered as a continuous covariate on the natural logarithm scale. "
        f"For log-transformed lymphocyte density, "
        f"{hr_str(cox['lymp_mean']['hr'], cox['lymp_mean']['ci_l'], cox['lymp_mean']['ci_u'], cox['lymp_mean']['p'])}. "
        f"For log-transformed tumour cell density, "
        f"{hr_str(cox['tumor_mean']['hr'], cox['tumor_mean']['ci_l'], cox['tumor_mean']['ci_u'], cox['tumor_mean']['p'])}. "
        f"For log-transformed tumour-to-lymphocyte ratio, "
        f"{hr_str(cox['ratio_mean']['hr'], cox['ratio_mean']['ci_l'], cox['ratio_mean']['ci_u'], cox['ratio_mean']['p'])}. "
        f"None of the three quantitative variables demonstrated a significant independent "
        f"association with overall survival in univariate analysis."
    ),
    (
        f"The absence of a survival signal from the digitally quantified variables in this "
        f"cohort likely reflects the moderate statistical power available with {n_ev} events, "
        f"the established strong prognostic dominance of pathological stage and nodal status "
        f"in early-stage lung adenocarcinoma, and the possibility that the prognostic effect "
        f"of tumour-infiltrating lymphocytes in this disease context is heterogeneous across "
        f"histological subtypes or confined to specific molecular subgroups not separately "
        f"analysed here. These considerations are addressed in the Discussion (Section 4.10)."
    ),
]

# ── 4.8 Multivariable Cox Models ──────────────────────────────────────────────
m1 = R['m1']
m2 = R['m2']
n_mv    = R['n_mv']
n_mv_ev = R['n_mv_events']
c1 = m1['c_index']
c2 = m2['c_index']

def mv_row(var, res, label):
    d = res[var]
    return (f"{label}: {hr_str(d['hr'], d['ci_l'], d['ci_u'], d['p'])}")

sec48_paras = [
    (
        f"Multivariable Cox proportional-hazards regression was performed to assess whether "
        f"digitally quantified morphological variables provide independent prognostic information "
        f"beyond established clinicopathological factors. Two nested models were constructed and "
        f"compared in {n_mv} cases with complete data ({n_mv_ev} death events)."
    ),
    (
        f"Model 1 (base model) incorporated the three conventional clinicopathological "
        f"prognostic factors identified in Chapter 2: pathological stage group, STAS, and "
        f"nodal status (pN). Results were as follows: "
        f"{mv_row('Stage_group', m1['results'], 'Stage group')}; "
        f"{mv_row('STAS', m1['results'], 'STAS')}; "
        f"{mv_row('pN', m1['results'], 'nodal status (pN)')}. "
        f"The base model achieved a concordance index (C-statistic) of {c1:.3f}. "
        f"Nodal status showed a trend toward independent prognostic significance "
        f"({fp(m1['results']['pN']['p'])})."
    ),
    (
        f"Model 2 (extended model) augmented the base model by adding per-case mean "
        f"lymphocyte density and per-case mean tumour cell density as continuous covariates "
        f"on the natural logarithm scale. Results were as follows: "
        f"{mv_row('Stage_group', m2['results'], 'Stage group')}; "
        f"{mv_row('STAS', m2['results'], 'STAS')}; "
        f"{mv_row('pN', m2['results'], 'nodal status (pN)')}; "
        f"{mv_row('log_lymp', m2['results'], 'log-lymphocyte density')}; "
        f"{mv_row('log_tumor', m2['results'], 'log-tumour cell density')}. "
        f"The extended model achieved a C-statistic of {c2:.3f}."
    ),
    (
        f"Comparison of the two models revealed that the addition of the two digital pathology "
        f"variables did not improve the discriminatory performance of the prognostic model "
        f"(C-statistic: {c1:.3f} for Model 1 versus {c2:.3f} for Model 2; "
        f"ΔC = {c2-c1:+.3f}). Neither log-transformed lymphocyte density "
        f"nor log-transformed tumour cell density reached statistical significance as "
        f"independent prognostic variables in the multivariable context ({fp(m2['results']['log_lymp']['p'])} "
        f"and {fp(m2['results']['log_tumor']['p'])}, respectively). "
        f"These results indicate that, in the present cohort, digitally quantified tumour "
        f"cell and lymphocyte densities do not independently predict overall survival beyond "
        f"the information provided by conventional clinicopathological staging variables."
    ),
    (
        f"The interpretation of these null findings is discussed in the context of study "
        f"power, follow-up duration, and biological heterogeneity in Section 4.10."
    ),
]

# ─── Open the docx and replace sections ───────────────────────────────────────
doc = Document('thesis_draft.docx')

def find_section_para_indices(doc, section_marker):
    """Find all paragraph indices for paragraphs starting with the section_marker."""
    indices = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(section_marker):
            indices.append(i)
    return indices

def get_run_font(para):
    """Get font settings from the first non-empty run of a paragraph."""
    for run in para.runs:
        if run.text.strip():
            return run.font.name, run.font.size, run.font.bold, run.font.italic
    return None, None, None, None

def replace_section_content(doc, start_marker, end_marker, new_paras):
    """
    Replace all paragraphs between start_marker (exclusive) and end_marker (exclusive)
    with new_paras text.
    """
    paras = doc.paragraphs
    start_idx = None
    end_idx   = None

    for i, p in enumerate(paras):
        if p.text.strip().startswith(start_marker) and start_idx is None:
            start_idx = i
        elif start_idx is not None and p.text.strip().startswith(end_marker):
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        print(f"WARNING: Could not find section boundaries: '{start_marker}' → '{end_marker}'")
        return

    # Get a reference paragraph for style (use first content paragraph after heading)
    ref_para = paras[start_idx + 1]
    ref_font_name, ref_font_size, _, _ = get_run_font(ref_para)

    # Delete paragraphs between start and end (exclusive)
    # We do this by clearing text and marking for removal
    # docx doesn't have a direct delete-paragraph API; we manipulate the XML
    from lxml import etree
    body = doc.element.body

    # Collect XML elements to delete
    all_body_children = list(body)
    para_elements = [p._element for p in paras]

    start_elem = para_elements[start_idx]
    end_elem   = para_elements[end_idx]

    # Find positions in body
    start_pos = list(body).index(start_elem)
    end_pos   = list(body).index(end_elem)

    # Build new paragraph XML elements
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    new_elements = []
    for text in new_paras:
        new_p = OxmlElement('w:p')
        # Copy paragraph properties from reference paragraph
        if ref_para._element.pPr is not None:
            import copy
            new_pPr = copy.deepcopy(ref_para._element.pPr)
            new_p.append(new_pPr)
        # Add run
        new_r = OxmlElement('w:r')
        # Copy run properties
        if ref_para.runs:
            ref_run = ref_para.runs[0]
            if ref_run._element.rPr is not None:
                import copy
                new_rPr = copy.deepcopy(ref_run._element.rPr)
                new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.text = text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        new_elements.append(new_p)

    # Remove old content paragraphs (between start+1 and end-1 inclusive)
    children = list(body)
    to_remove = children[start_pos + 1 : end_pos]
    for elem in to_remove:
        body.remove(elem)

    # Insert new elements after start_elem
    insert_after = start_elem
    for new_elem in reversed(new_elements):
        insert_after.addnext(new_elem)

    print(f"  Replaced {len(to_remove)} paragraphs with {len(new_elements)} new paragraphs "
          f"in section starting '{start_marker[:40]}'")


# Replace section 4.4
print("Updating section 4.4...")
replace_section_content(doc,
    '4.4  Quantitative Variables',
    '4.5  Intra-tumoral Variability',
    sec44_paras)

# Replace section 4.5
print("Updating section 4.5...")
replace_section_content(doc,
    '4.5  Intra-tumoral Variability',
    '4.6  Association with Clinicopathologic Features',
    sec45_paras)

# Replace section 4.6
print("Updating section 4.6...")
replace_section_content(doc,
    '4.6  Association with Clinicopathologic Features',
    '4.7  Survival Analysis',
    sec46_paras)

# Replace section 4.7
print("Updating section 4.7...")
replace_section_content(doc,
    '4.7  Survival Analysis',
    '4.8  Multivariable Prognostic Models',
    sec47_paras)

# Replace section 4.8
print("Updating section 4.8...")
replace_section_content(doc,
    '4.8  Multivariable Prognostic Models',
    '4.9  Figures',
    sec48_paras)

# Save
doc.save('thesis_final.docx')
print("\nSaved thesis_final.docx")
