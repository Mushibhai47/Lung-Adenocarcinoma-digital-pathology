"""
Appends Chapters 4 (sections 4.4-4.11), 5, and 6 to the existing thesis_draft.docx.
Uses real computed statistics from the CSV files.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── COMPUTED STATISTICS (from CSV analysis) ──────────────────────────────────
# Lymphocyte density (per-slide means, cells/mm2)
LYMP_N         = 171
LYMP_MEAN      = 442.0
LYMP_SD        = 691.5
LYMP_MEDIAN    = 116.5
LYMP_Q1        = 23.2
LYMP_Q3        = 580.9
LYMP_IQR       = 557.7
LYMP_MIN       = 0.0
LYMP_MAX       = 3478.8
LYMP_SKEW      = 2.335

# Tumor cell density (per-slide means, cells/mm2)
TUMOR_N        = 171
TUMOR_MEAN     = 1172.7
TUMOR_SD       = 819.6
TUMOR_MEDIAN   = 1132.9
TUMOR_Q1       = 433.1
TUMOR_Q3       = 1766.8
TUMOR_IQR      = 1333.7
TUMOR_MIN      = 4.0
TUMOR_MAX      = 3792.9
TUMOR_SKEW     = 0.520

# Intra-tumoral variability (n=167 slides with exactly 2 ROIs)
VARIAB_N       = 167
LYMP_R         = 0.868
LYMP_ICC       = 0.856
LYMP_MAD       = 213.19   # mean absolute difference cells/mm2
TUMOR_R        = 0.898
TUMOR_ICC      = 0.897
TUMOR_MAD      = 286.71

# Tumor-to-lymphocyte ratio (n=170, excluding 1 zero-lymphocyte slide)
TL_MEDIAN      = 6.37
TL_MEAN        = 22.43
TL_SD          = 46.67
TL_MIN         = 0.50
TL_MAX         = 329.14

# Tertile boundaries
TERTILE_LOW    = 39.8
TERTILE_HIGH   = 360.1

# ─── HELPER FUNCTIONS ─────────────────────────────────────────────────────────

def set_line_spacing(paragraph, spacing=1.5):
    paragraph.paragraph_format.line_spacing = spacing

def add_body_paragraph(doc, text, bold=False, italic=False):
    """Add a justified, Times New Roman 12pt, 1.5-spaced body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(12)
    run.bold       = bold
    run.italic     = italic
    return p

def add_chapter_heading(doc, text):
    """Bold, 14pt, 1.5-spaced left-aligned chapter heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing  = 1.5
    p.paragraph_format.space_before  = Pt(18)
    p.paragraph_format.space_after   = Pt(12)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(14)
    run.bold       = True
    return p

def add_section_heading(doc, text):
    """Bold, 13pt, 1.5-spaced left-aligned section heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing  = 1.5
    p.paragraph_format.space_before  = Pt(12)
    p.paragraph_format.space_after   = Pt(6)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(13)
    run.bold       = True
    return p

def add_subsection_heading(doc, text):
    """Bold, 12pt, 1.5-spaced left-aligned subsection heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing  = 1.5
    p.paragraph_format.space_before  = Pt(10)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(12)
    run.bold       = True
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_page_break())
    return p

def docx_page_break():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br


# ─── OPEN EXISTING DOCUMENT ───────────────────────────────────────────────────
import docx as _docx_module
doc = Document('C:/Users/musha/Desktop/QuPath/thesis_draft.docx')

# Add page break before new content
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run()
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — Sections 4.4 through 4.11
# ═══════════════════════════════════════════════════════════════════════════════

# ── 4.4 Quantitative Variables ─────────────────────────────────────────────────
add_section_heading(doc, '4.4  Quantitative Variables')

add_body_paragraph(doc,
    'The primary quantitative variables in this study were tumour cell density and lymphocyte '
    'density, both expressed as the number of detected cells per square millimetre (cells/mm²) '
    'of assessed tumour tissue. Tumour cell density was defined as the total count of nucleated '
    'cells detected within each rectangular region of interest (ROI) after subtracting '
    'lymphocyte-classified cells, divided by the measured area of that ROI in mm². Lymphocyte '
    'density was defined as the count of cells meeting morphological criteria consistent with '
    'lymphocytes (nuclear area 15–80 µm², circularity greater than 0.65, haematoxylin optical '
    'density above 0.40, and cell-to-nucleus area ratio below 2.5) divided by the ROI area. '
    'Both variables were computed independently for each ROI, and the per-case value reported '
    'for statistical analyses represents the arithmetic mean of the two ROI measurements '
    'obtained for each case, as described in section 4.3.'
)

add_body_paragraph(doc,
    'The tumour-to-lymphocyte ratio (T:L ratio) was derived as a secondary composite variable '
    'by dividing the per-case mean tumour cell density by the per-case mean lymphocyte density. '
    'This ratio provides a dimensionless index reflecting the relative balance between malignant '
    'epithelial cellularity and immune infiltration within the assessed tumour regions. One case '
    'exhibited a mean lymphocyte density of zero across both ROIs and was excluded from T:L ratio '
    'calculations to avoid division by zero; this case was retained in all other analyses.'
)

add_body_paragraph(doc,
    f'Quantitative measurements were successfully extracted from all {LYMP_N} cases processed '
    'through the digital analysis pipeline. For lymphocyte density, the distribution across '
    f'cases was markedly right-skewed (skewness = {LYMP_SKEW:.2f}), reflecting a pattern in which '
    'the majority of cases harboured moderate or low immune infiltration, while a minority '
    'exhibited substantially elevated lymphocyte counts. Specifically, the mean lymphocyte '
    f'density was {LYMP_MEAN:.1f} ± {LYMP_SD:.1f} cells/mm² (mean ± SD), with a median of '
    f'{LYMP_MEDIAN:.1f} cells/mm² (interquartile range [IQR]: {LYMP_Q1:.1f}–{LYMP_Q3:.1f} cells/mm²) '
    f'and a range of {LYMP_MIN:.1f}–{LYMP_MAX:.1f} cells/mm². The pronounced discrepancy between '
    'the mean and median values confirms substantial positive skew, consistent with the known '
    'heterogeneity of immune infiltration in lung adenocarcinoma. Tertile analysis yielded '
    f'boundaries at {TERTILE_LOW:.1f} and {TERTILE_HIGH:.1f} cells/mm², defining low (below '
    f'{TERTILE_LOW:.1f} cells/mm², n = 57, 33.3%), intermediate ({TERTILE_LOW:.1f}–{TERTILE_HIGH:.1f} '
    'cells/mm², n = 57, 33.3%), and high (above '
    f'{TERTILE_HIGH:.1f} cells/mm², n = 57, 33.3%) infiltration groups. The median-split threshold '
    f'of {LYMP_MEDIAN:.1f} cells/mm² was applied for dichotomous survival analyses.'
)

add_body_paragraph(doc,
    f'Tumour cell density was more symmetrically distributed (skewness = {TUMOR_SKEW:.2f}), with '
    f'a mean of {TUMOR_MEAN:.1f} ± {TUMOR_SD:.1f} cells/mm², a median of {TUMOR_MEDIAN:.1f} cells/mm² '
    f'(IQR: {TUMOR_Q1:.1f}–{TUMOR_Q3:.1f} cells/mm²), and a range spanning '
    f'{TUMOR_MIN:.1f}–{TUMOR_MAX:.1f} cells/mm². The broader IQR of {TUMOR_IQR:.1f} cells/mm² '
    'indicates substantial case-to-case variation in tumour cellularity, which likely reflects '
    'genuine biological heterogeneity in the degree of tumour compaction, stromal content, and '
    'necrosis across adenocarcinoma histological subtypes. The tumour-to-lymphocyte ratio, '
    f'computed for the {LYMP_N - 1} cases with non-zero lymphocyte density, had a median of '
    f'{TL_MEDIAN:.2f} (mean {TL_MEAN:.2f} ± {TL_SD:.2f}; range {TL_MIN:.2f}–{TL_MAX:.2f}), '
    'with a highly right-skewed distribution driven by a small number of cases with very low '
    'lymphocytic infiltration relative to tumour cellularity.'
)

# ── 4.5 Intra-tumoral Variability ──────────────────────────────────────────────
add_section_heading(doc, '4.5  Intra-tumoral Variability')

add_body_paragraph(doc,
    'Intratumoral heterogeneity is a fundamental challenge in tumour biology and pathological '
    'quantification. The spatial distribution of immune cells and tumour cells within a single '
    'neoplasm is non-uniform, reflecting localised differences in vascularity, necrosis, '
    'invasion front characteristics, and tumour-immune crosstalk. For any quantitative '
    'measurement derived from a limited sampling of a whole-slide image to be clinically '
    'meaningful, it must demonstrate acceptable reproducibility across sampling regions. '
    'This study addressed this issue prospectively by requiring two spatially separated '
    'rectangular ROIs per case, drawn to avoid necrotic areas, artefacts, and the tumour '
    'periphery where immune activity may differ substantially from the tumour core.'
)

add_body_paragraph(doc,
    f'Agreement between the two ROI measurements was assessed in the {VARIAB_N} cases for which '
    'exactly two ROIs were available. For lymphocyte density, the Pearson correlation coefficient '
    f'between ROI 1 and ROI 2 values was r = {LYMP_R:.3f} (p < 0.001), with an intraclass '
    f'correlation coefficient (ICC) of {LYMP_ICC:.3f}, indicating strong agreement between '
    'sampling regions. The mean absolute difference between paired ROI measurements for '
    f'lymphocyte density was {LYMP_MAD:.1f} cells/mm². For tumour cell density, the Pearson '
    f'correlation between ROI 1 and ROI 2 was r = {TUMOR_R:.3f} (p < 0.001), with an ICC of '
    f'{TUMOR_ICC:.3f} and a mean absolute difference of {TUMOR_MAD:.1f} cells/mm². Both correlation '
    'coefficients and ICC values exceeded the commonly applied threshold of 0.75 for good '
    'to excellent agreement, indicating that the two-ROI approach yields measurements that '
    'are substantially consistent within individual cases.'
)

add_body_paragraph(doc,
    'The observed correlations support the view that, despite genuine biological heterogeneity '
    'within tumours, the two-ROI sampling strategy captures sufficient information to characterise '
    'each case in a reproducible manner. The mean absolute differences—while numerically '
    'substantial in absolute terms, particularly for lymphocyte density—must be interpreted '
    'in the context of the wide intercase range of values. The within-case differences are '
    'considerably smaller than the intercase variability, ensuring that cases can be meaningfully '
    'discriminated using averaged two-ROI measurements. These findings reinforce the validity '
    'of the per-case mean as the primary analytical unit and are consistent with published '
    'evidence that two representative tumour regions provide adequate sampling for immune '
    'quantification in solid tumours of comparable size. It should be acknowledged, however, '
    'that two ROIs drawn from a single representative slide per case cannot fully capture '
    'three-dimensional tumour heterogeneity, and this remains a limitation of the present '
    'methodology.'
)

# ── 4.6 Association with Clinicopathologic Features ────────────────────────────
add_section_heading(doc, '4.6  Association with Clinicopathologic Features')

add_body_paragraph(doc,
    'Formal statistical analysis of associations between digitally quantified morphological '
    'variables and clinicopathologic features—including pathological stage, spread through '
    'air spaces (STAS) status, and nodal metastasis burden—requires completed case-level '
    'integration of scan identifiers with clinical records. This mapping process was ongoing '
    'at the time of the present analysis and will be finalised in the final version of this '
    'thesis. The following section describes the planned analytical approach and presents the '
    'anticipated structure of results, with specific numerical values reserved as placeholders '
    'to be populated upon data linkage.'
)

add_body_paragraph(doc,
    'The primary clinicopathologic variables of interest are: (i) STAS status (present versus '
    'absent), as identified by routine pathological examination and already established as an '
    'independent prognostic factor in the cohort described in Chapter 2; (ii) pathological '
    'stage group (stage I versus stage II–III, or stage IA versus IB–III), reflecting '
    'anatomical extent of disease at resection; and (iii) nodal metastasis status, specifically '
    'the number of involved nodal stations and the total nodal involvement ratio (NIR). '
    'Secondary variables include histological subtype (lepidic-predominant versus non-lepidic), '
    'visceral pleural invasion (VPI), and lymphovascular invasion (LVI).'
)

add_body_paragraph(doc,
    'Comparisons of lymphocyte density and tumour cell density across binary categorical variables '
    '(STAS, VPI, nodal positivity) will employ the Mann-Whitney U test, as both density variables '
    'are non-normally distributed (confirmed by Shapiro-Wilk testing). For ordinal variables '
    'with more than two categories (stage group, number of involved nodal stations), the '
    'Kruskal-Wallis test will be applied, with post-hoc pairwise comparisons using Dunn\'s '
    'procedure corrected for multiple comparisons via the Benjamini-Hochberg method. Correlation '
    'between continuous clinicopathologic variables (such as NIR) and digital density measures '
    'will be assessed using Spearman\'s rank correlation coefficient. All tests will be two-sided '
    'with an alpha threshold of 0.05.'
)

add_body_paragraph(doc,
    'Lymphocyte density showed [direction] association with STAS status (median [X] versus [Y] '
    'cells/mm² for STAS-positive versus STAS-negative tumours, Mann-Whitney U, p = [value]). '
    'Tumour cell density was [higher/lower] in STAS-positive cases compared to STAS-negative '
    'cases (median [X] versus [Y] cells/mm²; p = [value]). With respect to pathological stage, '
    'a [significant/non-significant] difference in lymphocyte density was observed across stage '
    'groups (Kruskal-Wallis H = [value], p = [value]), with post-hoc analysis suggesting '
    '[direction of pairwise differences]. Tumour cell density [did/did not] differ significantly '
    'by stage group (p = [value]). Spearman correlation between lymphocyte density and NIR '
    'was rₑ = [value] (p = [value]), and between tumour cell density and NIR was rₑ = [value] '
    '(p = [value]). These findings [support/do not support] the hypothesis that immune '
    'infiltration is modulated by the extent of regional dissemination in lung adenocarcinoma.'
)

add_body_paragraph(doc,
    'In addition to the primary analytical comparisons described above, exploratory subgroup '
    'analyses will examine whether digital density variables differ by histological predominant '
    'pattern (lepidic versus acinar versus papillary versus micropapillary/solid), given the '
    'established biological and prognostic differences between growth patterns. Lepidic-predominant '
    'tumours, characterised by low-grade biology and typically lower invasive potential, would '
    'be hypothesised to harbour different immune microenvironmental characteristics compared '
    'to solid or micropapillary-predominant tumours. Results of these exploratory analyses will '
    'be reported in the final integrated version of this chapter.'
)

# ── 4.7 Survival Analysis ──────────────────────────────────────────────────────
add_section_heading(doc, '4.7  Survival Analysis')

add_body_paragraph(doc,
    'The biological rationale for survival analysis of immune infiltration in non-small cell lung '
    'cancer (NSCLC) is firmly established in the literature. Tumour-infiltrating lymphocytes (TILs) '
    'have been shown to exert anti-tumour immune surveillance, and higher intratumoral lymphocyte '
    'density has been associated with improved overall survival (OS) and disease-free survival (DFS) '
    'in multiple NSCLC cohorts. Accordingly, the primary hypothesis of this component of the study '
    'is that higher lymphocyte density, as quantified by the digital analysis pipeline, is '
    'associated with longer overall survival after surgical resection of lung adenocarcinoma, '
    'independently of established prognostic factors including stage, STAS, and nodal status. '
    'A secondary hypothesis states that higher tumour cell density, reflecting greater tumour '
    'cellularity, is associated with shorter survival, consistent with the concept that densely '
    'cellular, aggressive tumours confer worse prognosis.'
)

add_body_paragraph(doc,
    'Overall survival was defined as the interval from the date of surgical resection to the '
    'date of death from any cause, with censoring at the date of last known follow-up for '
    'patients alive at the end of the observation period. Survival analyses will be performed '
    'following linkage of digital morphological data with clinical outcome data, which includes '
    'vital status and date of last contact for all cases in the cohort. Kaplan-Meier survival '
    'estimates will be generated for subgroups defined by dichotomised digital variables. '
    'Lymphocyte density will be dichotomised at the cohort median of '
    f'{LYMP_MEDIAN:.1f} cells/mm² into low-infiltration (below or equal to median) and '
    'high-infiltration (above median) groups. Tumour cell density will be similarly '
    f'dichotomised at its cohort median of {TUMOR_MEDIAN:.1f} cells/mm². Between-group '
    'survival differences will be assessed using the log-rank test.'
)

add_body_paragraph(doc,
    'Kaplan-Meier analysis of overall survival by lymphocyte infiltration group demonstrated '
    '[direction] survival in cases with high lymphocyte density compared to those with low '
    'lymphocyte density (median OS: [X] versus [Y] months; log-rank p = [value]). The '
    'five-year OS probability was [X]% in the high-infiltration group versus [Y]% in the '
    'low-infiltration group (hazard ratio [HR] = [value], 95% CI [lower–upper]; univariable '
    'Cox regression). For tumour cell density, survival was [direction] among cases with high '
    'tumour cell density relative to those with low density (median OS: [X] versus [Y] months; '
    'log-rank p = [value]; HR = [value], 95% CI [lower–upper]).'
)

add_body_paragraph(doc,
    'These results will be interpreted in the context of existing evidence, including the '
    'Immunoscore validated by Pagès and colleagues across multiple cancer types, the TIL '
    'scoring frameworks endorsed by the International TIL Working Group, and published '
    'computational pathology studies demonstrating the prognostic relevance of digitally '
    'quantified immune infiltration in NSCLC. Consistency with, or divergence from, this '
    'literature will be discussed with reference to the specific biological and methodological '
    'features of the present cohort.'
)

# ── 4.8 Multivariable Prognostic Models ────────────────────────────────────────
add_section_heading(doc, '4.8  Multivariable Prognostic Models')

add_body_paragraph(doc,
    'To determine whether digital morphological variables provide prognostic information '
    'independent of established clinicopathologic factors, a sequential multivariable modelling '
    'strategy will be employed using Cox proportional-hazards regression. Two nested models '
    'will be constructed and compared. The base model (Model 1) includes the clinicopathologic '
    'variables identified as significant in the analyses described in Chapter 2: pathological '
    'stage, STAS status, and nodal metastasis burden (categorised as absent, single-station, '
    'or multi-station involvement). This model establishes the baseline predictive performance '
    'of conventional pathological variables. Proportional-hazards assumptions will be assessed '
    'for each covariate using Schoenfeld residuals; variables violating this assumption will '
    'be entered as time-varying covariates or restricted to specific follow-up periods as '
    'appropriate.'
)

add_body_paragraph(doc,
    'The extended model (Model 2) augments Model 1 by adding lymphocyte density and tumour '
    'cell density as continuous covariates, entered on the natural logarithm scale to reduce '
    'the influence of extreme values and to linearise their relationships with the log-hazard. '
    'Continuous variables will be scaled by their interquartile range to produce hazard ratios '
    'interpretable per IQR increment, facilitating comparison of effect sizes between variables '
    'with different measurement scales. Model fit will be compared using the likelihood ratio '
    'test (comparing the log-likelihoods of Models 1 and 2), and model discrimination will '
    'be assessed using Harrell\'s concordance index (C-statistic) with bootstrapped 95% '
    'confidence intervals.'
)

add_body_paragraph(doc,
    'In Model 1, the clinicopathologic base model, STAS status (HR = [value], 95% CI '
    '[lower–upper]; p = [value]), advanced pathological stage (HR = [value], 95% CI '
    '[lower–upper]; p = [value]), and multi-station nodal involvement (HR = [value], 95% CI '
    '[lower–upper]; p = [value]) were [significant/non-significant] predictors of overall '
    'survival. The C-statistic for Model 1 was [value] (95% CI [lower–upper]). In the extended '
    'Model 2, lymphocyte density per IQR increment was associated with a HR of [value] (95% CI '
    '[lower–upper]; p = [value]), and tumour cell density per IQR increment was associated '
    'with a HR of [value] (95% CI [lower–upper]; p = [value]). Addition of the digital '
    'variables [significantly/did not significantly] improve model fit (likelihood ratio '
    'χ² = [value], df = 2, p = [value]), and the C-statistic increased to [value] '
    '(95% CI [lower–upper]).'
)

add_body_paragraph(doc,
    'If confirmed as statistically significant contributors in Model 2, these digital variables '
    'would support the concept that the immune microenvironment provides prognostic information '
    'that is not captured by standard pathological staging alone. This would be consistent with '
    'the growing body of evidence indicating that immune contextual features—the density, '
    'location, and functional orientation of infiltrating immune cells—modulate cancer '
    'outcomes independently of conventional disease extent. In practical terms, a significant '
    'and independent contribution of lymphocyte density would provide rationale for incorporating '
    'digitally quantified immune scores into prognostic reporting frameworks for surgically '
    'resected lung adenocarcinoma. Conversely, if digital variables do not add independent '
    'prognostic information beyond the base model, this may reflect limitations in the cell '
    'classification approach, the small number of ROIs per case, or the specific biological '
    'characteristics of this single-centre cohort.'
)

# ── 4.9 Figures ────────────────────────────────────────────────────────────────
add_section_heading(doc, '4.9  Figures')

add_body_paragraph(doc,
    'The following figure legends describe the planned illustrations for Chapter 4. Final '
    'figures will be generated upon completion of case-level data linkage and statistical '
    'analyses.'
)

figures = [
    ('Figure 4.1.', 'Representative haematoxylin and eosin (H&E) stained whole-slide image of '
     'surgically resected lung adenocarcinoma at 10× magnification, illustrating the '
     'characteristic histological features used to guide ROI selection, including the tumour '
     'parenchyma, stromal compartment, and areas of immune infiltration.'),
    ('Figure 4.2.', 'Annotated whole-slide image illustrating the two-ROI sampling strategy. '
     'Two rectangular regions of interest (each 4 mm²) are shown superimposed on the tumour '
     'area, separated by at least one high-power field, and positioned to represent tumour '
     'cellularity while avoiding necrosis, haemorrhage, and the tumour periphery.'),
    ('Figure 4.3.', 'Cell detection overlay for lymphocyte quantification in a representative '
     'case. Detected lymphocytes, meeting the morphological classification criteria (nuclear '
     'area 15–80 µm², circularity >0.65, haematoxylin optical density >0.40), are '
     'highlighted in green against the haematoxylin and eosin background. Non-lymphocyte '
     'cells are shown in blue.'),
    ('Figure 4.4.', 'Cell detection overlay for tumour cell quantification in the same '
     'representative case as Figure 4.3. Tumour cells (all detected cells minus lymphocyte-classified '
     'cells) are indicated in red. The relative spatial distribution of tumour cells and '
     'lymphocytes within the ROI is visible.'),
    ('Figure 4.5.', 'Scatter plot depicting the correlation between ROI 1 and ROI 2 '
     'lymphocyte density measurements across all cases with exactly two ROIs available '
     f'(n = {VARIAB_N}). Each point represents one case. The diagonal reference line '
     f'denotes perfect agreement. Pearson r = {LYMP_R:.3f} (p < 0.001). Axes are on a '
     'square-root scale for visualisation clarity.'),
    ('Figure 4.6.', 'Scatter plot depicting the correlation between ROI 1 and ROI 2 '
     'tumour cell density measurements across the same case series '
     f'(n = {VARIAB_N}). Pearson r = {TUMOR_R:.3f} (p < 0.001).'),
    ('Figure 4.7.', 'Box plots comparing lymphocyte density (cells/mm²) between STAS-positive '
     'and STAS-negative cases. Boxes represent the interquartile range (IQR), horizontal '
     'lines denote medians, whiskers extend to 1.5× IQR, and individual outliers are plotted. '
     'Statistical significance assessed by Mann-Whitney U test (p = [value]).'),
    ('Figure 4.8.', 'Box plots comparing tumour cell density (cells/mm²) across pathological '
     'stage groups (stage IA, IB, II, III). Statistical significance assessed by Kruskal-Wallis '
     'test with post-hoc Dunn\'s pairwise comparisons (overall p = [value]).'),
    ('Figure 4.9.', 'Kaplan-Meier overall survival curves stratified by lymphocyte density '
     f'(high > {LYMP_MEDIAN:.1f} cells/mm² versus low ≤ {LYMP_MEDIAN:.1f} cells/mm²). '
     'The number of patients at risk at selected time points is displayed below the x-axis. '
     'Log-rank p = [value].'),
    ('Figure 4.10.', 'Kaplan-Meier overall survival curves stratified by tumour cell density '
     f'(high > {TUMOR_MEDIAN:.1f} cells/mm² versus low ≤ {TUMOR_MEDIAN:.1f} cells/mm²). '
     'Log-rank p = [value].'),
    ('Figure 4.11.', 'Forest plot displaying adjusted hazard ratios and 95% confidence intervals '
     'from the multivariable Cox proportional-hazards extended model (Model 2), including '
     'pathological stage, STAS status, nodal involvement, lymphocyte density, and tumour cell '
     'density. Hazard ratios for continuous variables are expressed per IQR increment.'),
]

for fig_label, fig_text in figures:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run_label = p.add_run(fig_label + ' ')
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)
    run_label.bold = True
    run_text = p.add_run(fig_text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    run_text.bold = False

# ── 4.10 Discussion ────────────────────────────────────────────────────────────
add_section_heading(doc, '4.10  Discussion')

add_body_paragraph(doc,
    'This chapter presents a systematic digital pathology workflow for the quantification of '
    'tumour cell density and lymphocyte density in surgically resected lung adenocarcinoma, '
    'implemented entirely within the open-source QuPath platform. The analysis pipeline '
    'successfully extracted quantitative morphological data from all '
    f'{LYMP_N} cases processed, yielding a comprehensive dataset of paired density measurements '
    'for statistical analysis. The principal findings relate to the wide intercase variability '
    'in immune infiltration, the high intracase reproducibility of the two-ROI approach, '
    'and the marked positive skew in the lymphocyte density distribution—findings that '
    'are consistent with published literature on the immune microenvironment of NSCLC.'
)

add_body_paragraph(doc,
    'The prognostic relevance of tumour-infiltrating lymphocytes in NSCLC has been documented '
    'across multiple study designs and methodological approaches. Dønnem and colleagues '
    'demonstrated that stromal and intratumoral CD8-positive T-cell density, as assessed by '
    'immunohistochemistry, was independently associated with improved overall survival in '
    'resected NSCLC, with a dose-response relationship between immune cell density and '
    'survival advantage (Dønnem et al., 2011). Petersen and colleagues similarly found that '
    'high TIL density, as assessed on H&E slides using standardised visual scoring, was a '
    'significant positive prognostic factor in stage I–III NSCLC (Petersen et al., 2016). '
    'More recently, computational pathology approaches have been applied to replicate and '
    'extend these findings at scale; Saltz and colleagues demonstrated that deep learning-based '
    'spatial analysis of TILs across thousands of TCGA cases revealed consistent prognostic '
    'associations across multiple cancer types, including lung adenocarcinoma (Saltz et al., 2018). '
    'The present study contributes to this literature by applying computational quantification '
    'to a well-characterised single-centre surgical cohort, with the advantage of complete '
    'pathological re-evaluation and prospectively defined clinical variables.'
)

add_body_paragraph(doc,
    'A distinctive methodological feature of the present study is the use of a subtractive '
    'approach to operationally define tumour cell density: rather than training a deep learning '
    'classifier to positively identify tumour cells by morphology, tumour cell density is '
    'estimated as all detected nucleated cells minus those meeting lymphocyte classification '
    'criteria. This approach is computationally simple and does not require annotated training '
    'datasets for a tumour cell classifier, which represents a practical advantage in the '
    'setting of limited resources and heterogeneous staining quality. However, the approach '
    'carries an inherent limitation: it assumes that the detected non-lymphocyte cells are '
    'predominantly tumour cells, whereas in reality this population also includes stromal '
    'fibroblasts, endothelial cells, macrophages, and other inflammatory cells that do not '
    'meet the lymphocyte morphological criteria. Regions of dense stroma or prominent '
    'inflammatory infiltrate may therefore lead to overestimation of tumour cell density. '
    'This limitation was partially mitigated by the ROI selection strategy, which targeted '
    'areas of clear tumour cellularity as assessed by the reviewing pathologist, but cannot '
    'be fully resolved without a positive tumour cell classifier.'
)

add_body_paragraph(doc,
    'The intracase reproducibility findings are among the most methodologically important '
    'contributions of this chapter. The high Pearson correlation (r = '
    f'{LYMP_R:.3f} for lymphocyte density; r = {TUMOR_R:.3f} for tumour cell density) and '
    f'ICC values ({LYMP_ICC:.3f} and {TUMOR_ICC:.3f} respectively) between ROI 1 and ROI 2 '
    'measurements demonstrate that the two-ROI approach is substantially reproducible within '
    'cases, despite the known spatial heterogeneity of immune infiltration in solid tumours. '
    'These ICC values exceed the threshold of 0.75 conventionally considered to represent '
    'good to excellent agreement, providing empirical support for the validity of per-case '
    'averaged measurements as the primary analytical unit. The mean absolute differences '
    f'({LYMP_MAD:.1f} cells/mm² for lymphocyte density and {TUMOR_MAD:.1f} cells/mm² for '
    'tumour cell density), while numerically notable, are substantially smaller than the '
    'intercase variability (IQR of '
    f'{LYMP_IQR:.1f} and {TUMOR_IQR:.1f} cells/mm² respectively), confirming that case '
    'discrimination is preserved.'
)

add_body_paragraph(doc,
    'The QuPath-based workflow used in this study offers several practical advantages that '
    'merit discussion. As an open-source, platform-independent application, QuPath is freely '
    'accessible to any institution with adequate computational resources, and its scripting '
    'interface via Groovy allows complete automation of the analysis pipeline, ensuring '
    'methodological consistency across all cases. The reproducibility of threshold-based '
    'cell classification, however, is contingent on the consistency of staining across slides, '
    'which varied to some degree across the study cohort given that sections were obtained '
    'from archival paraffin blocks processed over an extended time period. Staining variability '
    'affects both haematoxylin optical density measurements and nuclear morphometry, '
    'potentially introducing systematic differences in classification across slides processed '
    'in different batches. Stain normalisation algorithms, such as the Macenko or Vahadane '
    'methods, could be applied as a preprocessing step in future iterations of this pipeline '
    'to reduce this source of variability.'
)

add_body_paragraph(doc,
    'Limitations of the present approach beyond those already discussed include the restriction '
    'to two ROIs per case, which may not capture the full extent of intratumoral heterogeneity '
    'in larger tumours; the use of rule-based morphological thresholds for cell classification, '
    'which are less accurate than trained convolutional neural network classifiers in '
    'distinguishing lymphocytes from histiocytes, plasma cells, and other mononuclear cells; '
    'the single-centre design, which may limit generalisability to cohorts with different '
    'staining protocols, scanner platforms, or patient populations; and the retrospective '
    'design, which carries inherent selection bias given that cases were restricted to those '
    'with available archival material and adequate tissue quality for digital analysis.'
)

add_body_paragraph(doc,
    'Future directions for this line of research include the implementation of deep '
    'learning-based cell classifiers—such as HoverNet or StarDist—which can simultaneously '
    'detect and classify multiple cell types with superior accuracy compared to threshold-based '
    'approaches. Integration of spatial analysis metrics, such as the proximity of lymphocytes '
    'to tumour cells or the calculation of nearest-neighbour distances between cell types, '
    'would add a further dimension to the characterisation of the tumour immune microenvironment. '
    'Correlation of digitally quantified lymphocyte density with PD-L1 immunohistochemistry '
    'scores and with tumour mutational burden, where available, would provide important '
    'biological validation of the computational measurements. Extension of the workflow to '
    'other NSCLC subtypes, particularly squamous cell carcinoma, and to other solid tumour '
    'types, would test the generalisability and scalability of the approach.'
)

# ── 4.11 Conclusion ────────────────────────────────────────────────────────────
add_section_heading(doc, '4.11  Conclusion')

add_body_paragraph(doc,
    'Chapter 4 demonstrates that a reproducible, scalable digital pathology workflow, '
    'implemented using open-source software and applied to routine haematoxylin and eosin-stained '
    'whole-slide images, can successfully extract quantitative morphological information from '
    'surgically resected lung adenocarcinoma without the need for additional staining, specialised '
    'reagents, or proprietary analysis platforms. The measurement of tumour cell density and '
    'lymphocyte density from two representative tumour regions per case yielded internally '
    'consistent, reproducible data across a cohort of '
    f'{LYMP_N} cases, with high intracase correlation between sampling regions (ICC '
    f'{LYMP_ICC:.3f} for lymphocyte density and {TUMOR_ICC:.3f} for tumour cell density) '
    'supporting the validity of the two-ROI approach. The marked heterogeneity in lymphocyte '
    f'density—ranging from {LYMP_MIN:.0f} to {LYMP_MAX:.0f} cells/mm², with a right-skewed '
    'distribution and a median of '
    f'{LYMP_MEDIAN:.1f} cells/mm²—reflects the genuine biological variability in '
    'immune infiltration across lung adenocarcinomas and underscores the importance of '
    'objective quantification for capturing differences that qualitative visual assessment '
    'cannot reliably discriminate. Pending completion of case-level clinical data linkage, '
    'these quantitative digital metrics represent a promising foundation for prognostic '
    'modelling that complements and potentially extends the conventional clinicopathologic '
    'staging framework.'
)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — Integrated Discussion (~2800 words)
# ═══════════════════════════════════════════════════════════════════════════════

# Page break before Chapter 5
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run()
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)

add_chapter_heading(doc, 'CHAPTER 5\nIntegrated Discussion')

add_body_paragraph(doc,
    'The practice of surgical pathology has undergone a profound transformation over recent '
    'decades, evolving from a discipline rooted in descriptive morphological observation toward '
    'an increasingly data-driven, quantitative science. This evolution has been driven by '
    'converging forces: advances in molecular diagnostics that have revealed biological '
    'heterogeneity invisible under the light microscope; recognition that qualitative, '
    'observer-dependent pathological assessments are subject to substantial interobserver '
    'variability; and the emergence of digital pathology platforms that enable systematic, '
    'reproducible extraction of quantitative features from tissue at scale. This thesis '
    'contributes to this trajectory by exploring three complementary dimensions through which '
    'prognostically relevant information can be extracted from routine haematoxylin and eosin-stained '
    'tissue sections of surgically resected lung adenocarcinoma: the systematic characterisation '
    'of established but sometimes underweighted clinicopathologic features, the recognition of '
    'environmental exposure information embedded within tissue, and the computational '
    'quantification of cellular composition of the tumour immune microenvironment. Together, '
    'these three dimensions illustrate a broader principle: that the routine diagnostic '
    'histological slide is a substantially underutilised data source, and that its systematic '
    'analysis—whether by trained pathologist recognition or by computational pipeline—can '
    'yield information with meaningful clinical and biological implications.'
)

add_section_heading(doc, '5.1  Conventional Prognostic Factors and Their Limitations')

add_body_paragraph(doc,
    'Chapter 2 of this thesis undertook a systematic clinicopathologic and molecular prognostic '
    'analysis of a single-centre cohort of surgically resected lung adenocarcinomas. The analysis '
    'confirmed that spread through air spaces (STAS), visceral pleural invasion (VPI), and '
    'nodal metastasis burden are independent adverse prognostic factors for overall survival '
    'following complete resection. These findings are concordant with the current evidence base '
    'and reinforce the importance of thorough pathological reporting of these features as part '
    'of the routine diagnostic workup. In particular, STAS—a relatively recently described '
    'pattern of tumour spread that encompasses the aerogenous migration of malignant cells '
    'beyond the main tumour mass into adjacent alveolar spaces—emerged as a significant '
    'predictor of outcomes in this cohort, consistent with multiple published series that have '
    'reported its association with increased recurrence risk and reduced survival across '
    'resection types and stages.'
)

add_body_paragraph(doc,
    'Despite the prognostic utility of these variables, their clinical application is not '
    'without limitations. Even within a single pathological stage, there is considerable '
    'heterogeneity in patient outcomes: stage I lung adenocarcinoma, which constitutes the '
    'largest and generally most favourable prognostic group, encompasses a broad spectrum of '
    'tumours with very different biological behaviours, ranging from indolent, lepidic-predominant '
    'lesions with near-normal life expectancy to compact, solid-predominant tumours with '
    'substantial recurrence risk despite complete resection. The current TNM staging system, '
    'while clinically validated and universally adopted, was primarily designed to reflect '
    'anatomical extent of disease rather than biological aggressiveness. Its prognostic '
    'resolution is therefore inherently limited in early-stage disease, where anatomical '
    'extent is by definition restricted. STAS and VPI partially address this limitation '
    'by incorporating histological invasion features, but both are assessed qualitatively '
    'and are subject to interobserver variability, particularly for STAS, where the boundary '
    'between genuine aerogenous spread and artefactual displacement of tumour cells during '
    'tissue processing has been a subject of ongoing debate. There is therefore a clear '
    'rationale for developing additional, more objective stratification tools that can be '
    'applied to the same routine tissue sections.'
)

add_section_heading(doc, '5.2  Under-recognised Histologic Information: Silica as a Paradigm')

add_body_paragraph(doc,
    'Chapter 3 of this thesis addressed a distinct but conceptually related question: whether '
    'routine haematoxylin and eosin-stained slides contain histologically identifiable evidence '
    'of environmental exposure that is not systematically reported as part of current diagnostic '
    'practice. The specific focus was on crystalline silica exposure, identified by the '
    'characteristic birefringent silica particles and associated granulomatous or fibrotic '
    'parenchymal changes visible under polarised light microscopy. The finding that histologic '
    'evidence of silica exposure was present in 12% of lobectomy specimens from the Central '
    'Portugal cohort—without this finding being reflected in clinical records or pathology '
    'reports in the majority of cases—provides a concrete illustration of the broader '
    'principle that slides contain more information than is systematically extracted through '
    'routine diagnostic examination.'
)

add_body_paragraph(doc,
    'This finding carries both direct clinical implications and broader conceptual significance. '
    'Directly, unrecognised silica exposure may be relevant to patient management decisions '
    'involving immunosuppressive therapy, occupational health reporting obligations, and '
    'surveillance for silicosis-related complications including accelerated fibrosis and '
    'increased susceptibility to infection. The geographic clustering of silica-positive cases '
    'within historically industrialised districts of Central Portugal—where quarrying, '
    'construction, and ceramics industries have historically employed large proportions of the '
    'working-age population—reflects the persistence of occupationally relevant environmental '
    'exposures in this region and underscores the importance of contextualising pathological '
    'findings within their epidemiological and geographic setting. At a broader conceptual '
    'level, the silica chapter exemplifies the principle that pathological slides encode '
    'information about a patient\'s biological and environmental history that extends well '
    'beyond the immediate diagnostic question. The systematic retrieval of this information '
    'requires appropriate methodology—in this case, the inclusion of polarised light '
    'examination as part of the routine lung pathology protocol—but does not require '
    'additional tissue sampling or specialised testing beyond what is already performed.'
)

add_body_paragraph(doc,
    'This paradigm of retrieving latent information from existing tissue connects directly '
    'to the overarching thesis of this work: that H&E histological slides are an underutilised '
    'data source whose systematic analysis—through both enhanced pathologist recognition '
    'of morphological features and computational quantification of cellular composition—can '
    'provide clinically and biologically relevant information that is currently not captured '
    'in routine diagnostic reporting. The silica chapter provides the qualitative case study, '
    'and the digital pathology chapter provides the quantitative operationalisation of this principle.'
)

add_section_heading(doc, '5.3  Quantitative Digital Morphology: Toward Objective Tumour Assessment')

add_body_paragraph(doc,
    'Chapter 4 of this thesis operationalised the quantitative dimension of the central thesis '
    'by implementing a computational image analysis pipeline to extract tumour cell density and '
    'lymphocyte density from two representative tumour regions per case across the full surgical '
    'cohort. The immune microenvironment of NSCLC has well-established prognostic relevance, '
    'with higher densities of tumour-infiltrating lymphocytes consistently associated with '
    'improved survival outcomes across published series. The innovation of the present approach '
    'lies not in the biological hypothesis—which is well-supported by prior work—but in '
    'its methodological implementation: a fully open-source, script-driven, threshold-based '
    'pipeline that can be applied systematically to large archival cohorts without the need '
    'for immunohistochemistry, specialised staining, or proprietary software licences.'
)

add_body_paragraph(doc,
    'The quantification of immune infiltration in clinical pathology has traditionally relied '
    'on qualitative visual estimation, which is known to be subject to interobserver variability '
    'and is difficult to standardise across institutions. Scoring frameworks such as the '
    'Immunoscore and the International TIL Working Group\'s H&E-based scoring system have '
    'provided structured approaches to quantitative immune assessment, but both still rely '
    'on trained human observers performing either enumeration or scoring of defined cell '
    'populations. Computational quantification removes this subjectivity, providing continuous '
    'numerical outputs that can be reproduced exactly from the same image, are not subject '
    'to fatigue or anchoring effects, and can be applied at scale across hundreds of cases '
    'in a fraction of the time required for manual assessment. The reproducibility of the '
    'two-ROI approach demonstrated in Chapter 4—with Pearson r = '
    f'{LYMP_R:.3f} and ICC = {LYMP_ICC:.3f} for lymphocyte density, and r = {TUMOR_R:.3f} and '
    f'ICC = {TUMOR_ICC:.3f} for tumour cell density, both within the range conventionally '
    'considered indicative of good to excellent agreement—provides empirical support for '
    'the internal consistency of this approach.'
)

add_body_paragraph(doc,
    'The two-ROI strategy was specifically designed to balance practical feasibility with '
    'adequate sampling of intratumoral heterogeneity. Requiring exhaustive whole-slide '
    'cell quantification, while theoretically ideal, is computationally expensive and time-consuming '
    'at scale. The demonstrated high intracase correlation between spatially separated ROIs '
    'suggests that two adequately selected representative regions are sufficient to characterise '
    'the predominant immune context of each tumour, and that the per-case averaged measurement '
    'provides a reliable summary statistic. This does not imply that immune heterogeneity '
    'within tumours is absent—the mean absolute differences between ROIs '
    f'({LYMP_MAD:.1f} cells/mm² for lymphocyte density and {TUMOR_MAD:.1f} cells/mm² for '
    'tumour cell density) are non-trivial—but rather that the intercase variation is '
    'substantially greater than the intracase variation, preserving the discriminatory utility '
    'of the per-case average.'
)

add_section_heading(doc, '5.4  Toward Integrated Prognostic Models')

add_body_paragraph(doc,
    'The underlying motivation for combining conventional clinicopathologic analysis (Chapter 2) '
    'with digital morphological quantification (Chapter 4) is the hypothesis that these two '
    'information sources are complementary rather than redundant. Conventional staging captures '
    'the anatomical extent of disease and the presence of specific invasion patterns, but does '
    'not directly measure the biological competence of the host immune response, which is '
    'increasingly recognised as a determinant of tumour behaviour and treatment sensitivity. '
    'Digital quantification of immune infiltration operationalises this host response component '
    'in a numerical, reproducible form that can be entered into multivariable statistical models.'
)

add_body_paragraph(doc,
    'The extended Cox regression model described in Chapter 4, which adds lymphocyte density '
    'and tumour cell density to the conventional base model, directly tests whether digital '
    'morphological metrics provide prognostic information independent of stage, STAS, and '
    'nodal status. This is the most rigorous available test of whether digital quantification '
    'adds genuine incremental value beyond established pathological variables. The analogy '
    'from adjacent fields is instructive: in breast cancer pathology, the prognostic utility '
    'of TILs as assessed on H&E has been established even after adjustment for established '
    'clinicopathologic variables including nodal status, tumour grade, and molecular subtype, '
    'suggesting that immune infiltration captures biological information orthogonal to these '
    'conventional variables. If a similar pattern is observed in the present lung adenocarcinoma '
    'cohort, it would lend support to the development of integrated prognostic scores that '
    'combine conventional staging with digital immune metrics.'
)

add_body_paragraph(doc,
    'The tumour-to-lymphocyte ratio computed in this study—with a median of '
    f'{TL_MEDIAN:.2f} across cases with non-zero lymphocyte density, and a highly right-skewed '
    'distribution reflecting the predominance of tumour-rich, immune-poor phenotypes in a '
    'substantial proportion of cases—provides a single integrated metric that captures '
    'the relative balance between malignant cellularity and immune surveillance. This ratio '
    'has conceptual parallels with the neutrophil-to-lymphocyte ratio (NLR) used in blood-based '
    'systemic inflammation indices, which has demonstrated prognostic relevance in multiple '
    'cancer types. The digital T:L ratio operationalises an analogous concept at the local '
    'intratumoral level, potentially providing superior biological specificity for tumour '
    'microenvironmental phenotyping. Whether this ratio adds independent prognostic information '
    'beyond either density variable alone will be examined in the multivariable analyses '
    'described in Chapter 4.'
)

add_section_heading(doc, '5.5  Translational Implications')

add_body_paragraph(doc,
    'The translational context of this work is shaped by the rapid maturation of digital '
    'pathology as a clinical technology. Whole-slide imaging systems have received regulatory '
    'approval for primary diagnosis in multiple jurisdictions, including FDA clearance for '
    'primary digital pathology diagnosis in the United States, and adoption of digital workflows '
    'is accelerating across European pathology departments. As this infrastructure becomes '
    'routine, the computational analysis of digitised slides—performed alongside or '
    'subsequent to primary diagnostic reporting—becomes an increasingly practical proposition '
    'rather than a purely research activity.'
)

add_body_paragraph(doc,
    'The QuPath-based workflow described in this thesis is specifically designed to be compatible '
    'with this emerging clinical infrastructure. QuPath operates on standard whole-slide image '
    'formats without requiring proprietary software or hardware, can be driven by fully '
    'automated scripted pipelines, and produces structured numerical outputs suitable for '
    'integration into laboratory information systems. The complete automation of the analysis '
    'pipeline—from image import through ROI selection guidance, cell detection, classification, '
    'and density calculation to results export—means that, in principle, the workflow could '
    'be applied to new cases with minimal additional pathologist time, potentially as a '
    'background computational process triggered at the time of slide digitisation. The '
    'incremental cost of extracting digital morphological metrics from a slide that has already '
    'been digitised for primary diagnosis is therefore minimal.'
)

add_body_paragraph(doc,
    'The integration of digitally quantified immune metrics into clinical prognostic reporting '
    'would, in practice, require prospective validation in independent cohorts, standardisation '
    'of ROI selection criteria to ensure reproducibility across pathologists and institutions, '
    'and establishment of clinical decision thresholds informed by survival data. The present '
    'study represents an important foundational step toward this goal by demonstrating '
    'feasibility, internal reproducibility, and analytical validity in a moderately sized '
    'retrospective cohort. It does not, however, constitute clinical validation, and the '
    'translation of these findings into routine practice would require the further studies '
    'outlined in section 5.7.'
)

add_section_heading(doc, '5.6  Limitations')

add_body_paragraph(doc,
    'The limitations of this thesis must be considered carefully when interpreting the findings '
    'and extrapolating to other settings. First, all three study components draw from a single '
    'institutional cohort from the Centro Hospitalar Universitário de Coimbra (CHUC), a tertiary '
    'referral centre in Central Portugal. While this single-centre design ensures uniformity '
    'in surgical technique, pathological handling, and diagnostic standards, it restricts '
    'generalisability to populations with different epidemiological profiles, staining protocols, '
    'or case-mix characteristics. In particular, the silica exposure findings reflect the '
    'specific occupational history of the Central Portugal region and may not be representative '
    'of cohorts from other geographic areas.'
)

add_body_paragraph(doc,
    'Second, the retrospective design introduces the possibility of selection bias, as cases '
    'were restricted to those with available archival material of adequate quality for digital '
    'analysis, and survival data require linkage to administrative records that may be '
    'incomplete for patients lost to follow-up. Third, the digital analysis pipeline relies '
    'on threshold-based morphological classification, which is sensitive to staining variability '
    'across archival sections processed at different times. Batch-to-batch variation in '
    'haematoxylin intensity and tissue processing can affect the optical density measurements '
    'and nuclear morphometry that underpin cell classification, potentially introducing '
    'systematic differences in measured density across slide batches. Fourth, the restriction '
    'to two ROIs per case limits the capture of spatial immune heterogeneity across the full '
    'tumour section, and cases with only one or more than two ROIs (a minority in this dataset) '
    'were handled differently in reproducibility analyses, potentially introducing slight '
    'inconsistency in per-case measurement accuracy.'
)

add_body_paragraph(doc,
    'Fifth, the case-level linkage between slide identifiers and clinical records required '
    'additional validation steps that were not fully complete at the time of this analysis, '
    'meaning that survival analyses and clinicopathologic associations—described in detail '
    'in Chapter 4—could not be reported with final numerical values in this thesis version. '
    'The presentation of these sections as structured analyses with placeholder values for '
    'specific results, rather than as completed analyses, reflects this practical constraint '
    'rather than a conceptual gap in the study design. Sixth, and most substantially from a '
    'methodological standpoint, the rule-based lymphocyte classifier used in this study is '
    'considerably less accurate than modern deep learning approaches for cell-type classification. '
    'Histiocytes, plasma cells, and some inflammatory cells share morphological features with '
    'lymphocytes at the threshold parameters applied, and misclassification of these cell types '
    'in both directions (lymphocyte overcounting and undercounting) is possible and cannot be '
    'fully quantified without manual validation in a representative subsample.'
)

add_section_heading(doc, '5.7  Future Directions')

add_body_paragraph(doc,
    'The findings and limitations of this thesis define a clear agenda for future research. '
    'The most immediately impactful methodological improvement would be the replacement of '
    'the threshold-based lymphocyte classifier with a trained deep learning cell segmentation '
    'and classification model. Tools such as HoverNet, which performs simultaneous nuclear '
    'instance segmentation and cell-type classification using a multi-branch architecture, '
    'and StarDist, which employs star-convex polygon representations for robust nuclear '
    'detection in dense tissue, have demonstrated superior cell-type classification accuracy '
    'compared to threshold-based approaches in multiple published benchmarks. Retaining '
    'the QuPath framework but replacing the classification component with a HoverNet or '
    'StarDist-based model would substantially improve the specificity of lymphocyte detection '
    'and would allow positive identification of tumour cells as a distinct class, eliminating '
    'the methodological limitations of the subtractive approach used here.'
)

add_body_paragraph(doc,
    'Beyond improved cell classification, the integration of spatial analysis metrics represents '
    'a scientifically compelling extension of the present work. The proximity of lymphocytes '
    'to tumour cells—quantified, for example, as the mean nearest-neighbour distance between '
    'classified lymphocytes and tumour cells—may carry additional prognostic information '
    'beyond simple density measures, as it captures the degree of immune-tumour interface '
    'formation that is hypothesised to underlie effective anti-tumour immune surveillance. '
    'Graph-based spatial analysis methods, including those implemented in toolkits such as '
    'squidpy and Giotto, provide accessible approaches for computing such metrics from '
    'cell coordinate outputs of the QuPath pipeline.'
)

add_body_paragraph(doc,
    'Multi-centre validation of the digital quantification workflow is essential before clinical '
    'translation can be considered. A validation cohort from a geographically and institutionally '
    'distinct centre, processed with a standardised ROI selection protocol and the same '
    'analytical pipeline, would allow assessment of whether the density measurements and their '
    'prognostic associations are reproducible across different staining protocols, scanner '
    'platforms, and patient populations. The development and application of stain normalisation '
    'as a preprocessing step would facilitate this cross-institutional reproducibility. '
    'Integration with PD-L1 immunohistochemistry data—available for many cases in the cohort '
    'given its clinical use as a predictive biomarker for immunotherapy—would allow direct '
    'biological validation of the digital lymphocyte density measurements against an established '
    'tissue-based immune marker, providing a cross-platform correlation that could support '
    'the clinical interpretability of the digital metric. Finally, application of the complete '
    'integrated analytical framework—conventional prognostic factors, silica assessment, '
    'and digital immune quantification—to other thoracic tumour types, including squamous '
    'cell carcinoma and large cell neuroendocrine carcinoma, would test the generalisability '
    'of the conceptual and methodological framework developed in this thesis.'
)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — Final Conclusions (~600 words)
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run()
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)

add_chapter_heading(doc, 'CHAPTER 6\nConclusions')

add_body_paragraph(doc,
    'In a single-centre surgical cohort of 173 patients undergoing complete resection of '
    'lung adenocarcinoma at a tertiary referral centre in Central Portugal, systematic '
    'clinicopathologic and molecular analysis confirmed that spread through air spaces, '
    'visceral pleural invasion, and nodal metastasis burden are independent adverse prognostic '
    'factors for overall survival following resection, reinforcing the importance of thorough '
    'pathological reporting of these histological features as integral components of the '
    'diagnostic and prognostic workup. These findings are concordant with the contemporary '
    'literature and provide institutional validation of the prognostic relevance of STAS '
    'and nodal staging granularity beyond the dichotomous TNM nodal category, supporting '
    'their systematic inclusion in surgical pathology reports for resected lung adenocarcinoma.'
)

add_body_paragraph(doc,
    'Histological examination of the resected specimens using polarised light microscopy '
    'revealed evidence of crystalline silica exposure in 12% of lobectomy cases. In the '
    'majority of affected cases, this finding had not been previously documented in clinical '
    'records or pathology reports, indicating that environmentally relevant information '
    'embedded in routine H&E-stained sections remains largely unreported under current '
    'diagnostic practice. The geographic clustering of silica-positive cases within '
    'historically industrial districts of the Central Portugal region reflects the persistence '
    'of occupational silica exposure in this population and highlights the potential clinical '
    'and occupational health significance of systematically incorporating polarised light '
    'examination into the standard pathological assessment of lung resection specimens. '
    'This finding exemplifies the broader principle that routine diagnostic slides encode '
    'information extending beyond the primary diagnostic question, and that systematic '
    'application of appropriate methodology can retrieve this information without additional '
    'invasive procedures.'
)

add_body_paragraph(doc,
    'A reproducible digital pathology workflow was implemented using the open-source QuPath '
    'platform to quantify tumour cell density and lymphocyte density in two representative '
    f'tumour regions per case across all {LYMP_N} cases processed. The workflow demonstrated '
    f'strong intracase reproducibility, with an intraclass correlation coefficient of '
    f'{LYMP_ICC:.3f} for lymphocyte density and {TUMOR_ICC:.3f} for tumour cell density '
    'between paired ROI measurements, supporting the validity of the two-ROI approach as '
    'a practical and sufficiently reproducible sampling strategy for large-cohort computational '
    'morphology studies. The marked heterogeneity in lymphocyte density across cases—with '
    f'a median of {LYMP_MEDIAN:.1f} cells/mm², an IQR of {LYMP_IQR:.1f} cells/mm², and a '
    f'range spanning from {LYMP_MIN:.0f} to {LYMP_MAX:.0f} cells/mm²—reflects the genuine '
    'biological variability in tumour immune microenvironmental composition across '
    'lung adenocarcinomas and underscores the necessity of objective, quantitative methods '
    'for capturing this variability in a clinically usable form.'
)

add_body_paragraph(doc,
    'Quantitative digital metrics encompassing lymphocyte density and tumour cell density '
    'were successfully extracted from routine haematoxylin and eosin-stained whole-slide '
    'images without the requirement for additional immunohistochemical staining, specialised '
    'reagents, or proprietary analysis software, demonstrating the practical feasibility '
    'of large-scale computational morphology in a routine pathology setting equipped with '
    'standard whole-slide imaging infrastructure. The complete automation of the analysis '
    'pipeline, from cell detection through classification and density calculation to '
    'structured data export, ensures methodological consistency across cases and minimises '
    'the per-case analytical burden, making this approach applicable in principle to '
    'prospective clinical implementation.'
)

add_body_paragraph(doc,
    'The integration of digitally quantified morphological variables into multivariable '
    'prognostic models, incorporating conventional clinicopathologic factors established in '
    'Chapter 2 alongside lymphocyte density and tumour cell density from Chapter 4, offers '
    'the potential to improve prognostic stratification of surgically resected lung '
    'adenocarcinoma beyond what is achievable with anatomical staging and histological '
    'invasion features alone. The completion of case-level clinical data linkage for '
    'survival analysis—currently pending at the time of this submission—will enable '
    'definitive testing of whether digital immune quantification provides independent '
    'prognostic information, and the results will be incorporated into the final version '
    'of this thesis.'
)

add_body_paragraph(doc,
    'Collectively, these findings articulate and provide empirical support for a unifying '
    'thesis: that routine haematoxylin and eosin histological slides of surgically resected '
    'lung adenocarcinoma contain quantifiable morphological information that is currently '
    'underutilised in clinical practice. The systematic extraction of this information—through '
    'enhanced pathologist recognition of environmental and microenvironmental features and '
    'through computational quantification of cellular composition—represents a methodologically '
    'accessible and clinically meaningful step toward more objective, reproducible, and '
    'data-driven surgical pathology practice. As digital pathology infrastructure continues '
    'to mature and computational tools become increasingly accessible, the analytical framework '
    'demonstrated in this thesis provides a template for augmenting the prognostic information '
    'derived from the diagnostic tissue already available for every surgically resected '
    'lung cancer patient.'
)


# ─── SAVE ─────────────────────────────────────────────────────────────────────
doc.save('C:/Users/musha/Desktop/QuPath/thesis_draft.docx')
print('Done. thesis_draft.docx saved successfully.')
print(f'Total paragraphs in final document: {len(doc.paragraphs)}')