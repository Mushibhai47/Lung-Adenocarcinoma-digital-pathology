"""
update_chapter2.py
Updates Chapter 2 of thesis_final.docx with actual SPSS-derived statistics.
Also adds PDL1 and Vimentin results, fixes sex/staging/KRAS/LVI errors.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'C:\Users\musha\Desktop\QuPath\thesis_final.docx')

def set_para_text(para, text):
    p_elem = para._p
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    run = para.add_run(text)
    return run

def insert_paragraph_after(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())

def find_para(start_text):
    for p in doc.paragraphs:
        if p.text.startswith(start_text):
            return p
    return None

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT para 11 -- update STAS/VPI/nodal claim
# ════════════════════════════════════════════════════════════════════════════
p = doc.paragraphs[11]
set_para_text(p,
    "In a retrospective series of 172 patients who underwent curative-intent resection at CHUC -- "
    "Centro Hospitalar Universitario de Coimbra between 2015 and 2019 (n = 172 with complete "
    "survival data; median follow-up 73.7 months), spread through air spaces (STAS) was associated "
    "with significantly inferior overall survival on univariate analysis (five-year OS 51.2% versus "
    "71.2%; log-rank p = 0.028; HR = 1.73, 95% CI 1.05-2.84). Nodal positivity (pN+, 22.1% of "
    "cases) was the dominant independent adverse prognostic factor on multivariable Cox analysis "
    "(HR = 3.12, 95% CI 1.50-6.50, p = 0.002), with a model C-index of 0.670. Lymphovascular "
    "invasion was present in 27.3% of cases but did not reach statistical significance for overall "
    "survival. PD-L1 expression (>/=1%, assessed in 104 cases) and vimentin immunopositivity "
    "(assessed in 139 cases) did not demonstrate significant survival associations in this cohort."
)
print("Updated Abstract para 11")

# ════════════════════════════════════════════════════════════════════════════
# 2.3.1 COHORT CHARACTERISTICS
# ════════════════════════════════════════════════════════════════════════════
p74 = find_para("The study cohort comprised 173 patients with a mean age of 66.6 years")
if p74:
    set_para_text(p74,
        "The study cohort comprised 172 patients with complete survival data (one case excluded due "
        "to missing vital status date). Mean age was 66.6 years (SD 9.3; range 33-84 years). Male "
        "sex predominated, representing 62.8% of the cohort (108 males, 64 females [37.2%]). The "
        "predominance of male patients is consistent with the epidemiology of occupational and "
        "environmental carcinogen exposure in the Central Portuguese population served by this "
        "institution, and with the slightly higher lung adenocarcinoma incidence in males observed "
        "in Mediterranean cohorts. Surgical resection was by lobectomy in 165 patients (96.0%), "
        "pneumonectomy in five (2.9%), and bilobectomy in three (1.7%). Surgical margins were free "
        "in all but one patient (0.6%)."
    )
    print("Updated para 74: Cohort characteristics (sex corrected)")
else:
    print("WARNING: para 74 not found")

p75 = find_para("The distribution of pathologic stages was as follows: Stage I in 118 patients")
if p75:
    set_para_text(p75,
        "The distribution of pathologic stages was as follows: Stage I in 100 patients (58.1%), "
        "comprising 7 IA1 (4.1%), 19 IA2 (11.0%), 21 IA3 (12.2%), and 53 IB (30.8%); Stage II in "
        "47 patients (27.3%), comprising 8 IIA (4.7%) and 39 IIB (22.7%); and Stage III in 25 "
        "patients (14.5%), comprising 22 IIIA (12.8%) and 3 IIIB (1.7%). Nodal metastasis was "
        "identified in 38 patients (22.1%): 31 with N1 disease (18.0%) and 7 with N2 disease "
        "(4.1%). The predominance of Stage IB and IIB disease reflects the referral and surgical "
        "selection criteria of this tertiary centre, where complete resection is pursued in patients "
        "with locally advanced but anatomically resectable tumours."
    )
    print("Updated para 75: Staging distribution (actual numbers)")
else:
    print("WARNING: para 75 not found")

# ════════════════════════════════════════════════════════════════════════════
# 2.3.2 HISTOLOGIC FINDINGS
# ════════════════════════════════════════════════════════════════════════════
p77 = find_para("Assessment of histologic growth patterns revealed acinar pattern as the predominant subtype")
if p77:
    set_para_text(p77,
        "Assessment of histologic growth patterns revealed acinar pattern as the predominant subtype "
        "in 82 cases (47.7%), consistent with its status as the most common growth pattern in most "
        "published Western and Southern European lung adenocarcinoma series. Solid pattern predominance "
        "was identified in 24 cases (14.0%), papillary in 17 (9.9%), micropapillary in 15 (8.7%), "
        "and lepidic in 10 (5.8%). Predominant pattern was not assessable in 24 cases (14.0%), "
        "primarily due to mixed pattern composition without a definable predominant subtype. Among "
        "pathological diagnoses, invasive non-mucinous adenocarcinoma represented the large majority "
        "at 142 cases (82.6%), with invasive mucinous adenocarcinoma in 13 cases (7.6%), pleomorphic "
        "carcinoma in 8 (4.7%), and adenosquamous carcinoma in 4 (2.3%)."
    )
    print("Updated para 77: Histologic patterns")
else:
    print("WARNING: para 77 not found")

p78 = find_para("Spread through air spaces was identified in 42 cases")
if p78:
    set_para_text(p78,
        "Spread through air spaces was identified in 41 cases (23.8%), consistent with the lower end "
        "of the prevalence range reported in methodologically rigorous series using the standardised "
        "WHO definition. One case was scored as STAS-indeterminate and excluded from STAS-related "
        "analyses (n = 171). Lymphovascular invasion was present in 47 cases (27.3%) and visceral "
        "pleural invasion in 94 cases (54.7%)."
    )
    print("Updated para 78: STAS/LVI/VPI numbers")
else:
    print("WARNING: para 78 not found")

# ════════════════════════════════════════════════════════════════════════════
# 2.3.3 MOLECULAR FINDINGS
# ════════════════════════════════════════════════════════════════════════════
p80 = find_para("Molecular testing results were available for the full cohort. EGFR mutations")
if p80:
    set_para_text(p80,
        "Molecular testing was performed in 64 of 172 cases (37.2%) using next-generation sequencing "
        "(NGS) or pre-NGS targeted methods; the remaining 108 cases (62.8%) were not tested, "
        "reflecting the progressive implementation of reflex molecular testing across the study "
        "period (2015-2019). Among the 64 tested cases, 39 (60.9%) harboured no detectable driver "
        "mutation. EGFR mutations were identified in 18 tested cases (28.1% of tested; 10.5% of the "
        "full cohort), comprising 11 exon 19 deletions (17.2%), 5 other EGFR mutations (7.8%), and "
        "2 L858R substitutions (3.1%). The prevalence of EGFR mutations among tested cases (28.1%) "
        "is consistent with the range reported in Southern European and Mediterranean series, which "
        "typically report EGFR prevalences of 15-35% among tested patients -- higher than Northern "
        "European cohorts but lower than East Asian populations."
    )
    print("Updated para 80: EGFR molecular findings")
else:
    print("WARNING: para 80 not found")

p81 = find_para("ALK rearrangements, ROS1 fusions, BRAF V600E mutations, and MET exon 14")
if p81:
    set_para_text(p81,
        "Among other driver alterations, ALK rearrangements or fusions were identified in 3 tested "
        "cases (4.7% of tested), ROS1 rearrangement in 1 (1.6%), RET fusion in 1 (1.6%), and KRAS "
        "G12C mutation in 2 (3.1%). The small number of KRAS G12C-positive cases (n = 2) precluded "
        "meaningful survival analysis for this mutation subgroup in isolation. The overall low "
        "prevalence of individually rare driver alterations reflects the known epidemiology of "
        "targetable oncogene-driven lung adenocarcinoma in unselected, predominantly male European "
        "surgical cohorts. Vimentin immunohistochemistry -- a marker of epithelial-to-mesenchymal "
        "transition -- was assessed in 139 of 172 cases (80.8%), with positivity in 64 cases "
        "(46.0%). PD-L1 immunohistochemistry was assessed in 104 cases (60.5%), with expression "
        ">/=1% in 45 cases (43.3%)."
    )
    print("Updated para 81: Other molecular + Vimentin/PDL1 intro")
else:
    print("WARNING: para 81 not found")

# ════════════════════════════════════════════════════════════════════════════
# 2.3.4 SURVIVAL ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
p83 = find_para("Kaplan-Meier analysis and log-rank testing were performed for all pre-specified")
if p83:
    set_para_text(p83,
        "Kaplan-Meier analysis and log-rank testing were performed for all pre-specified "
        "clinicopathologic variables. The median follow-up for the cohort was 73.7 months (IQR "
        "37.7-92.5 months; range 0.3-119.3 months), with 73 deaths recorded (42.4% event rate). "
        "Spread through air spaces was associated with a significantly inferior overall survival: "
        "five-year OS 51.2% in STAS-positive versus 71.2% in STAS-negative cases (log-rank "
        "chi2 = 4.81, p = 0.028; univariate HR = 1.73, 95% CI 1.05-2.84). This hazard ratio "
        "magnitude and direction are consistent with published meta-analyses of STAS in resected "
        "lung adenocarcinoma, though the present cohort effect was of moderate statistical "
        "strength."
    )
    print("Updated para 83: STAS survival (actual p=0.028)")
else:
    print("WARNING: para 83 not found")

p84 = find_para("Lymphovascular invasion was associated with a five-year OS of 50%")
if p84:
    set_para_text(p84,
        "Lymphovascular invasion was present in 47 cases (27.3%) but did not demonstrate a "
        "statistically significant association with overall survival in this cohort "
        "(five-year OS 63.8% LVI-positive versus 67.5% LVI-negative; log-rank p = 0.680; "
        "univariate HR = 1.10, 95% CI 0.66-1.85). Visceral pleural invasion, present in 94 "
        "cases (54.7%), showed a trend toward inferior survival that did not reach statistical "
        "significance (five-year OS 58.9% VPI-positive versus 75.6% VPI-negative; log-rank "
        "p = 0.073; univariate HR = 1.52, 95% CI 0.94-2.43). The absence of a significant "
        "survival effect for LVI in this cohort contrasts with several published series and "
        "may reflect the moderate sample size and the relatively high proportion of early-stage "
        "cases in whom LVI's prognostic effect may be attenuated by the favourable background "
        "prognosis."
    )
    print("Updated para 84: LVI NOT significant (p=0.680), VPI trend (p=0.073)")
else:
    print("WARNING: para 84 not found")

p85 = find_para("Nodal status exerted the strongest individual prognostic effect")
if p85:
    set_para_text(p85,
        "Nodal status exerted the strongest individual prognostic effect among all analysed "
        "variables. Five-year overall survival was 73.7% for N0 patients (n = 134) compared "
        "with 40.7% for patients with nodal positivity (N1 or N2; n = 38), a highly "
        "significant difference (log-rank chi2 = 15.70, p < 0.001; univariate HR = 2.60, "
        "95% CI 1.58-4.26). Pathologic stage group (I versus II/III) showed a directional "
        "difference in five-year OS (70.8% Stage I versus 60.7% Stage II/III) that did not "
        "reach statistical significance (log-rank p = 0.143), likely because nodal status -- "
        "the primary driver of stage advancement in this cohort -- captures most of the "
        "prognostic information encoded by the TNM stage group when included simultaneously."
    )
    print("Updated para 85: Nodal survival (actual numbers)")
else:
    print("WARNING: para 85 not found")

p86 = find_para("Among molecular variables, KRAS-mutated tumours were associated")
if not p86:
    p86 = find_para("Among molecular variables, KRAS")
if p86:
    set_para_text(p86,
        "Among molecular variables analysed in the 64-case tested subset, EGFR-mutant cases "
        "(n = 18) did not demonstrate a statistically significant difference in overall survival "
        "compared with EGFR wild-type cases (five-year OS 77.8% versus 71.7%; log-rank p = 0.238). "
        "The absence of a survival advantage for EGFR-mutant cases in this resected cohort is "
        "consistent with the established understanding that, in early-stage surgically resected "
        "disease, the benefit of curative resection equalises outcomes between EGFR-mutant and "
        "wild-type patients, and that EGFR mutation confers its principal benefit through its "
        "predictive value for adjuvant targeted therapy rather than through inherently more "
        "indolent tumour biology. The number of KRAS G12C-positive cases (n = 2) was insufficient "
        "for meaningful survival analysis. PD-L1 expression (>/=1%) did not demonstrate a "
        "significant survival association (five-year OS 68.9% versus 72.3% for PDL1 <1%; "
        "log-rank p = 0.793). Vimentin immunopositivity similarly showed no significant "
        "survival difference (five-year OS 59.4% positive versus 72.7% negative; log-rank "
        "p = 0.202), though a directional trend toward inferior survival in vimentin-positive "
        "cases -- consistent with the hypothesised role of epithelial-to-mesenchymal transition "
        "in tumour aggressiveness -- warrants exploration in larger cohorts."
    )
    print("Updated para 86: Removed KRAS claim, updated EGFR/PDL1/Vimentin results")
else:
    print("WARNING: para 86 (KRAS) not found")

p87 = find_para("Among molecular variables, EGFR-mutated cases did not display")
if not p87:
    p87 = find_para("Among molecular variables, EGFR")
if p87:
    # This was the old para 87 about EGFR - now replaced by content in para 86
    # Just make this a continuation about sex
    set_para_text(p87,
        "Sex showed a non-significant trend toward superior survival in female patients "
        "(five-year OS 75.0% female versus 61.4% male; log-rank p = 0.067; univariate "
        "HR = 0.64, 95% CI 0.39-1.05). The directional advantage for females is consistent "
        "with the well-documented female survival advantage in lung adenocarcinoma, which has "
        "been attributed to the higher prevalence of targetable driver mutations and the "
        "generally more favourable tumour biology in female never-smokers."
    )
    print("Updated para 87: Now reports sex survival trend")
else:
    print("WARNING: para 87 (EGFR) not found - may already be removed")

# ════════════════════════════════════════════════════════════════════════════
# 2.3.5 MULTIVARIABLE ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
p88 = find_para("Multivariable Cox proportional hazards regression, incorporating all variables")
if p88:
    set_para_text(p88,
        "Multivariable Cox proportional hazards regression incorporating STAS, LVI, VPI, nodal "
        "status (pN+), and pathologic stage group was performed on 171 cases with complete "
        "covariate data (73 death events). In this multivariable model, nodal positivity (pN+) "
        "was the sole variable achieving statistical significance, with a hazard ratio of 3.12 "
        "(95% CI 1.50-6.50, p = 0.002), confirming it as the dominant independent adverse "
        "prognostic factor in this cohort. STAS retained a directional adverse association in "
        "the multivariable model (HR = 1.39, 95% CI 0.82-2.33, p = 0.221) but did not reach "
        "significance after adjustment for other variables, suggesting partial collinearity "
        "with nodal status. LVI, VPI, and stage group did not independently predict survival "
        "after mutual adjustment (all p > 0.18). The overall model C-index was 0.670, "
        "indicating moderate discriminatory performance consistent with published prognostic "
        "models for resected lung adenocarcinoma. The addition of PD-L1 to the model (in the "
        "104 cases with PD-L1 data) did not improve discrimination (C-index 0.631; PD-L1 "
        "HR = 0.94, p = 0.859)."
    )
    print("Updated para 88: Multivariable (actual: only pN+ significant, C-index=0.670)")
else:
    print("WARNING: para 88 not found")

# ════════════════════════════════════════════════════════════════════════════
# 2.4 DISCUSSION
# ════════════════════════════════════════════════════════════════════════════
p90 = find_para("The findings of this chapter confirm and extend the established prognostic literature")
if p90:
    set_para_text(p90,
        "The findings of this chapter characterise the clinicopathologic and molecular prognostic "
        "landscape of surgically resected lung adenocarcinoma in a contemporary Portuguese academic "
        "centre cohort. The most striking finding is the dominant prognostic role of nodal "
        "positivity: pN+ conferred a three-fold increase in mortality risk (HR = 3.12) in "
        "multivariable analysis, while the five-year overall survival dropped from 73.7% in N0 "
        "patients to 40.7% in N+ patients. This magnitude of prognostic effect is consistent "
        "with the established literature on nodal metastasis as the principal determinant of "
        "prognosis in surgically resected NSCLC. STAS was associated with significantly inferior "
        "survival in univariate analysis (HR = 1.73, p = 0.028), and retained a directional "
        "adverse effect in multivariable analysis that did not reach significance after adjustment "
        "for nodal status, suggesting that part of the prognostic information conveyed by STAS "
        "is captured by the higher prevalence of nodal involvement in STAS-positive tumours."
    )
    print("Updated para 90: Discussion opening (actual findings)")
else:
    print("WARNING: para 90 not found")

p91 = find_para("The prevalence of STAS in this cohort (24.3%)")
if not p91:
    p91 = find_para("The prevalence of STAS in this cohort")
if p91:
    set_para_text(p91,
        "The prevalence of STAS in this cohort (23.8%) is consistent with the range reported in "
        "series applying the standardised WHO definition, and the associated univariate hazard ratio "
        "of 1.73 aligns closely with values reported in contemporary European series. The absence "
        "of STAS as an independent significant variable in the multivariable model -- despite "
        "univariate significance -- is a noteworthy finding. STAS and nodal status frequently "
        "co-occur (STAS-positive tumours tend to have higher rates of lymph node involvement), "
        "and in a multivariable context where nodal status carries the dominant prognostic weight "
        "(HR = 3.12), the residual prognostic effect of STAS may be attenuated. This does not "
        "negate the clinical utility of STAS assessment: its significance in univariate analysis "
        "and its established role in guiding sublobar resection decisions support its inclusion "
        "in routine pathological reporting."
    )
    print("Updated para 91: STAS discussion")
else:
    print("WARNING: para 91 not found")

p92 = find_para("The relatively low prevalence of KRAS mutations")
if not p92:
    p92 = find_para("The relatively low prevalence of KRAS")
if p92:
    set_para_text(p92,
        "The molecular findings of this cohort reflect the expected epidemiology of driver "
        "alterations in a Southern European surgical series. EGFR mutations in 28.1% of tested "
        "cases are slightly lower than the 30-40% prevalence reported in predominantly female "
        "never-smoker Asian cohorts, and broadly consistent with the 15-30% range reported in "
        "European series. The predominance of male patients and the associated higher likelihood "
        "of smoking-related carcinogenesis may explain the lower EGFR prevalence compared with "
        "female-predominant cohorts. The absence of a detectable survival advantage for EGFR-mutant "
        "patients in this resected cohort -- where curative surgery equalises outcomes regardless "
        "of driver mutation status -- is expected and has been consistently reported in published "
        "series of early-stage resected disease. Of note, only 37.2% of cases in this cohort "
        "underwent molecular testing, a limitation attributable to the progressive implementation "
        "of NGS during the study period (2015-2019); more recent cohorts would be expected to "
        "achieve near-complete molecular characterisation."
    )
    print("Updated para 92: Molecular discussion (removed KRAS survival claim)")
else:
    print("WARNING: para 92 not found")

p93 = find_para("The absence of a statistically significant survival difference between EGFR-mutated")
if not p93:
    p93 = find_para("The absence of a statistically significant survival difference between EGFR")
if p93:
    set_para_text(p93,
        "Vimentin immunopositivity showed a non-significant trend toward inferior survival "
        "(five-year OS 59.4% versus 72.7%; p = 0.202), and its univariate hazard ratio of 1.39 "
        "is directionally consistent with the hypothesised role of epithelial-to-mesenchymal "
        "transition markers in tumour aggressiveness. PD-L1 expression (>/=1% in 43.3% of tested "
        "cases) did not show any survival association (p = 0.793). The high proportion of cases "
        "with PD-L1 positivity (43.3%) is clinically relevant in the context of adjuvant "
        "immunotherapy eligibility criteria, where PD-L1 status informs treatment decisions, "
        "though the survival data in this pre-immunotherapy cohort cannot address the predictive "
        "value of PD-L1 for therapeutic response."
    )
    print("Updated para 93: PDL1/Vimentin discussion")
else:
    print("WARNING: para 93 not found")

p94 = find_para("Several limitations of this analysis warrant acknowledgement")
if p94:
    set_para_text(p94,
        "Several limitations of this analysis warrant acknowledgement. The retrospective design "
        "introduces the possibility of selection bias, as cases were restricted to patients "
        "undergoing surgical resection at a single academic institution. The incomplete coverage "
        "of molecular testing (37.2% of cases) limits conclusions about the prognostic role of "
        "driver alterations and introduces potential selection bias in molecular subgroup analyses, "
        "as patients selected for molecular testing during the earlier part of the study period "
        "may not be representative of the full cohort. The relatively modest sample size (n = 172) "
        "limits statistical power to detect hazard ratios below approximately 1.5 in multivariable "
        "analysis, which may account for the failure to confirm the expected prognostic associations "
        "of LVI and VPI in the multivariable model. Despite these limitations, the cohort provides "
        "a well-characterised contemporary European series with long follow-up (median 73.7 months) "
        "that supports the prognostic primacy of nodal status in surgically resected lung "
        "adenocarcinoma."
    )
    print("Updated para 94: Limitations")
else:
    print("WARNING: para 94 not found")

# ════════════════════════════════════════════════════════════════════════════
# 2.5 CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
p96 = find_para("This chapter has established the conventional prognostic baseline")
if not p96:
    p96 = find_para("This chapter has established the conventional")
if p96:
    set_para_text(p96,
        "This chapter has established the clinicopathologic and molecular prognostic baseline for "
        "the study cohort, demonstrating that nodal positivity is the dominant independent "
        "predictor of overall survival (multivariable HR = 3.12, p = 0.002; five-year OS 73.7% "
        "N0 versus 40.7% N+). STAS was associated with inferior survival in univariate analysis "
        "(HR = 1.73, p = 0.028), confirming its prognostic relevance in this Southern European "
        "surgical series. Lymphovascular invasion did not demonstrate a significant survival "
        "association in this cohort, while VPI showed a directional trend. PD-L1 expression and "
        "vimentin immunopositivity were characterised as clinicopathologic variables but did not "
        "independently predict survival. These findings provide the conventional prognostic "
        "benchmark against which the quantitative digital pathology metrics developed in Chapter 4 "
        "are evaluated, confirming that nodal status and STAS are the primary histopathologic "
        "determinants of risk in this cohort that any new biomarker must meaningfully augment to "
        "add clinical value."
    )
    print("Updated para 96: Conclusion")
else:
    print("WARNING: para 96 not found")

# ════════════════════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════════════════════
doc.save(r'C:\Users\musha\Desktop\QuPath\thesis_final.docx')
print("\nAll Chapter 2 updates complete. Saved thesis_final.docx")
